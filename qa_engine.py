"""
qa_engine.py — Core Source Document → HTML QA comparison engine

Goals:
- Compare source document visible text against HTML visible text token-by-token.
- Catch changed words, missing words, extra words, punctuation changes.
- Catch missing bold/strong tags where source text/style indicates bold where available.
- Catch missing heading tags where source text/style indicates a heading where available.
- Catch visible text outside expected semantic tags.
- Ignore Mytonomy footer/copyright boilerplate.
- Ignore DOCTYPE/head/script/style/meta/link/title.
- Ignore logo.png and blank/header logo placeholders.
- Compare non-logo images by perceptual hash.
"""

import os
QA_ENGINE_BUILD = "RTL-FIX-2026-08-23-v2"  # <-- confirms this exact file is loaded; check server startup log / print(QA_ENGINE_BUILD)
import re
import math
import io
import base64
import difflib
import unicodedata
import regex as uregex
from pathlib import Path
from dataclasses import dataclass

import pdfplumber
from bs4 import BeautifulSoup, NavigableString
from bs4.element import Comment, Doctype, Declaration, ProcessingInstruction
from PIL import Image
import numpy as np
import imagehash
from html.parser import HTMLParser


# ── Regex / constants ────────────────────────────────────────────────────────

WS = re.compile(r"\s+")
TOKEN_RE = uregex.compile(
    r"""
    [\p{L}\p{M}\p{N}]+(?:[-'’‐‒–—−][\p{L}\p{M}\p{N}]+)*
    |
    [\p{P}\p{S}]
    """,
    uregex.VERBOSE,
)

UNICODE_PUNCT_OR_SYMBOL_RE = uregex.compile(r"^[\p{P}\p{S}]$")
ARABIC_SCRIPT_RE = uregex.compile(r"\p{Script=Arabic}")
ARABIC_DIACRITICS_RE = uregex.compile(r"[ؐ-ًؚ-ٰٟۖ-ۭ]")
RTL_CONTROL_RE = uregex.compile(r"[‎‏‪-‮⁦-⁩]")
ZERO_WIDTH_RE = uregex.compile(r"[​‌‍﻿]")

MOJIBAKE = re.compile(r"Ã.|â€™|â€œ|â€\x9d|\ufffd")

PLACEHOLDER = re.compile(
    r"(lorem ipsum|\btbd\b|\btodo\b|insert (text|image|video) here"
    r"|\[placeholder\]|\bxxx\b|<<[^>]*>>)",
    re.IGNORECASE,
)

EXPECTED_BULLET_MARKER = "•"

BOILERPLATE_PATTERNS = [
    # PDF footer:
    # © 2026 Mytonomy, Inc. All Rights Reserved. Page 1
    r"(?:©\s*)?20\d{2}\s+Mytonomy,\s*Inc\.?\s+All\s+Rights\s+Reserved\.?(?:\s+Page\s+\d+)?",

    # HTML footer:
    # © 2026 Mytonomy, Inc. All rights reserved.
    r"(?:©\s*)?20\d{2}\s+Mytonomy,\s*Inc\.?\s+All\s+rights\s+reserved\.?",

    # Bare copyright line with no "All Rights Reserved" attached (seen in
    # localized PDF footers where that phrase is translated separately) --
    # e.g. "© 2025 Mytonomy, Inc." on its own.
    r"©\s*20\d{2}\s+Mytonomy\s*,?\s*Inc\.?",

    # Common generated page labels
    r"\bPage\s+\d+\b",

    # Localized copyright/footer boilerplate in the supplied RTF/HTML packages.
    r"تمام\s*حقوق\s*محفوظ\s*است[\s\.,،؛؟]*",
    r"جملہ\s*حقوق\s*محفوظ\s*ہیں[\s\.,،؛؟]*",
    r"جميع\s*الحقوق\s*محفوظة[\s\.,،؛؟]*",

    # Localized "Page N" label (Farsi: صفحه N) -- block order in the PDF's
    # extracted footer can come out either way ("صفحه 1" or "1 صفحه"),
    # verified directly against real source files.
    r"(?:صفحه\s*\d+|\d+\s*صفحه)",
    # Localized "Page N" label (Modern Standard Arabic: الصفحة N), same
    # either-order block issue as the Farsi pattern above.
    r"(?:الصفحة\s*\d+|\d+\s*الصفحة)",
]

IGNORE_IMAGES = {
    "logo.png",
    "company_logo.png",
    "brand_logo.png",
}

_IMG_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg"}

SKIP_TEXT_PARENT_TAGS = {
    "script",
    "style",
    "head",
    "title",
    "meta",
    "link",
    "noscript",
    "template",
}

ALLOWED_TEXT_BLOCKS = {
    "p",
    "li",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "td",
    "th",
    "figcaption",
    "caption",
    "blockquote",
    "dt",
    "dd",
}

LANG_TRIGRAMS = {
    "French": ["es ", "de ", "ent", "les", "on ", "la ", "ons", "que", "tion", "le "],
    "German": ["en ", "der", "ie ", "die", "und", "ein", "ich", "das", "sch", "ng "],
    "Spanish": ["de ", "es ", "en ", "que", "la ", "el ", "ón ", "ion", "os ", "con"],
    "Portuguese": ["de ", "os ", "as ", "que", "em ", "do ", "da ", "ão ", "es ", "ar "],
    "Italian": ["del", "che", "la ", "di ", "le ", "in ", "ne ", "on ", "per", "una"],
    "Dutch": ["en ", "de ", "het", "van", "ing", "ge ", "een", "te ", "dat", "ij "],
    "English": ["the", " th", "he ", "ing", "and", "ion", "ed ", "ent", "or ", "al "],
}

HEADER_LOGO_MAX_TOP_RATIO = 0.14
HEADER_LOGO_MAX_LEFT_RATIO = 0.22
HEADER_LOGO_MAX_HEIGHT_RATIO = 0.15
HEADER_LOGO_MAX_WIDTH_RATIO = 0.45

VOID_TAGS = {
    "area", "base", "br", "col", "embed", "hr", "img", "input",
    "link", "meta", "param", "source", "track", "wbr"
}

HYPHEN_TOKENS = {"-", "‐", "-", "‒", "–", "—"}

HYPHEN_CHARS_FOR_QA = "-‐-‒–—−"

OPTIONAL_CLOSE_TAGS = set()  # keep strict because your HTML is XHTML-like


# ── Dataclasses ──────────────────────────────────────────────────────────────

@dataclass
class Issue:
    category: str
    severity: str
    line: object
    message: str
    snippet: str = ""
    expected: str = ""
    actual: str = ""


@dataclass
class TextToken:
    text: str
    norm: str
    line: object = None
    tag: object = None
    bold: bool = False
    italic: bool = False
    heading_level: int = 0
    in_allowed_block: bool = True

class RawHTMLTagValidator(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=False)
        self.stack = []
        self.issues = []

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()

        if tag in VOID_TAGS:
            return

        self.stack.append((tag, self.getpos()[0]))

    def handle_startendtag(self, tag, attrs):
        # Handles <br />, <img />, etc.
        return

    def handle_endtag(self, tag):
        tag = tag.lower()
        line = self.getpos()[0]

        if tag in VOID_TAGS:
            return

        if not self.stack:
            self.issues.append(Issue(
                "Unexpected Closing Tag",
                "error",
                line,
                f'Found closing tag </{tag}> without a matching opening tag.',
                expected=f"<{tag}> before </{tag}>",
                actual=f"</{tag}>",
            ))
            return

        top_tag, top_line = self.stack[-1]

        if top_tag == tag:
            self.stack.pop()
            return

        # Search deeper in stack to see if this tag exists but nesting is wrong.
        matching_index = None
        for i in range(len(self.stack) - 1, -1, -1):
            if self.stack[i][0] == tag:
                matching_index = i
                break

        if matching_index is None:
            self.issues.append(Issue(
                "Unexpected Closing Tag",
                "error",
                line,
                f'Found closing tag </{tag}> but current open tag is <{top_tag}> from line {top_line}.',
                expected=f"</{top_tag}>",
                actual=f"</{tag}>",
            ))
            return

        # Example: <p><strong>Text</p></strong>
        unclosed = self.stack[matching_index + 1:]

        for unclosed_tag, unclosed_line in reversed(unclosed):
            self.issues.append(Issue(
                "Tag Nesting Mismatch",
                "error",
                line,
                f'Tag <{unclosed_tag}> opened on line {unclosed_line} was not closed before </{tag}>.',
                expected=f"</{unclosed_tag}> before </{tag}>",
                actual=f"</{tag}>",
            ))

        # Remove matched tag and wrongly nested children from stack.
        self.stack = self.stack[:matching_index]

    def close(self):
        super().close()

        for tag, line in reversed(self.stack):
            if tag in OPTIONAL_CLOSE_TAGS:
                continue

            self.issues.append(Issue(
                "Unclosed Tag",
                "error",
                line,
                f'Opening tag <{tag}> on line {line} was not closed.',
                expected=f"</{tag}>",
                actual="missing closing tag",
            ))


def check_raw_html_tags(html_path: str):
    with open(html_path, "r", encoding="utf-8", errors="replace") as f:
        raw_html = f.read()

    validator = RawHTMLTagValidator()

    try:
        validator.feed(raw_html)
        validator.close()
    except Exception as e:
        validator.issues.append(Issue(
            "HTML Parse Error",
            "error",
            None,
            f"HTML parser failed while checking raw tags: {e}",
            expected="Valid HTML",
            actual="Parse failure",
        ))

    return validator.issues

# ── Basic text helpers ───────────────────────────────────────────────────────

def normalize_for_token_scan(text: str) -> str:
    """
    Light Unicode normalization before tokenization.

    Keeps visible token text close to the original document while removing
    invisible/control characters and normalizing Unicode representation.
    """
    if text is None:
        return ""

    text = str(text)
    text = unicodedata.normalize("NFC", text)

    text = RTL_CONTROL_RE.sub("", text)
    text = ZERO_WIDTH_RE.sub("", text)
    text = text.replace(" ", " ")

    text = text.replace("“", '"').replace("”", '"')
    text = text.replace("‘", "'").replace("’", "'")

    text = text.replace("‐", "-")
    text = text.replace("‒", "-")
    text = text.replace("–", "-")
    text = text.replace("—", "-")
    text = text.replace("−", "-")

    text = uregex.sub(r"\s+", " ", text).strip()
    return text


def normalize_for_compare(text: str) -> str:
    """
    Strong comparison normalization for English, accented Latin languages,
    Arabic, Urdu, and mixed-language text.
    """
    if text is None:
        return ""

    text = str(text)
    text = unicodedata.normalize("NFC", text)

    text = RTL_CONTROL_RE.sub("", text)
    text = ZERO_WIDTH_RE.sub("", text)

    # Arabic/Urdu optional marks and elongation.
    text = ARABIC_DIACRITICS_RE.sub("", text)
    text = text.replace("ـ", "")

    text = text.replace(" ", " ")

    text = text.replace("“", '"').replace("”", '"')
    text = text.replace("‘", "'").replace("’", "'")

    text = text.replace("‐", "-")
    text = text.replace("‒", "-")
    text = text.replace("–", "-")
    text = text.replace("—", "-")
    text = text.replace("−", "-")

    # Arabic alef variants.
    text = text.replace("أ", "ا")
    text = text.replace("إ", "ا")
    text = text.replace("آ", "ا")
    text = text.replace("ٱ", "ا")

    # Arabic/Persian/Urdu kaf and yeh variants.
    text = text.replace("ك", "ک")
    text = text.replace("ي", "ی")
    text = text.replace("ى", "ی")

    # Light heh normalization.
    text = text.replace("ۀ", "ہ")

    # Arabic/Urdu punctuation variants.
    text = text.replace("،", ",")
    text = text.replace("؛", ";")
    text = text.replace("؟", "?")
    text = text.replace("۔", ".")

    text = uregex.sub(r"\s+", " ", text).strip()
    return text.lower()


def norm_text(text: str) -> str:
    return normalize_for_compare(text).rstrip(".")


def norm_token(text: str) -> str:
    return normalize_for_compare(text)


def strip_boilerplate_text(text: str) -> str:
    text = normalize_for_token_scan(text)

    for pattern in BOILERPLATE_PATTERNS:
        text = re.sub(pattern, " ", text, flags=re.IGNORECASE)

    return WS.sub(" ", text).strip()


def is_boilerplate_line(text: str) -> bool:
    raw = WS.sub(" ", text or "").strip()
    if not raw:
        return True

    cleaned = strip_boilerplate_text(raw)
    if not cleaned:
        return True

    # After stripping known boilerplate phrases, a line with nothing left
    # but stray punctuation (".", "..", ". .") is still boilerplate --
    # verified directly against real source files where the footer left a
    # dangling period once the copyright/rights/page-number phrases were
    # removed, which the emptiness check above didn't catch.
    if not re.search(r"[^\W\d_]", cleaned, flags=re.UNICODE):
        return True

    if re.fullmatch(r"Page\s+\d+", raw, flags=re.IGNORECASE):
        return True

    return False


def detect_language(text: str) -> str:
    """Lightweight language/script guess used for reporting only."""
    sample = normalize_for_token_scan(text[:3000]).lower()

    if len(re.findall(r"[Ѐ-ӿ]", sample)) > 20:
        return "Russian"

    arabic_script_count = len(ARABIC_SCRIPT_RE.findall(sample))
    if arabic_script_count > 20:
        urdu_specific_count = len(re.findall(r"[ٹڈڑگچپژکںھہےۓ]", sample))
        if urdu_specific_count >= 3:
            return "Urdu"
        return "Arabic"

    if len(re.findall(r"[ऀ-ॿ]", sample)) > 20:
        return "Hindi"
    if len(re.findall(r"[一-鿿]", sample)) > 20:
        return "Chinese"
    if len(re.findall(r"[぀-ヿ]", sample)) > 20:
        return "Japanese"
    if len(re.findall(r"[가-힯]", sample)) > 20:
        return "Korean"

    scores = {
        lang: sum(sample.count(t) for t in tgrams)
        for lang, tgrams in LANG_TRIGRAMS.items()
    }

    latin_diacritics = len(
        re.findall(
            r"[àáâãäåæçèéêëìíîïðñòóôõöøùúûüýÿœ]",
            sample,
        )
    )

    if latin_diacritics < 5:
        scores["English"] = scores.get("English", 0) * 1.5

    return max(scores, key=scores.get)


def context(tokens, idx, radius=8) -> str:
    if not tokens:
        return ""

    idx = max(0, min(idx, len(tokens) - 1))
    start = max(0, idx - radius)
    end = min(len(tokens), idx + radius + 1)

    return " ".join(t.text for t in tokens[start:end])


def is_punctuation(token: str) -> bool:
    return bool(UNICODE_PUNCT_OR_SYMBOL_RE.fullmatch(token or ""))


def has_case(text: str) -> bool:
    """True for scripts with upper/lowercase, false for Arabic/Urdu."""
    return any(ch.lower() != ch.upper() for ch in text or "")


def should_ignore_token(token: str) -> bool:
    return token in {"•", "●", "◦", "▪", "▫", "·", "‣", "⁃", "◉", "○", "■"}


# ── HTML extraction ──────────────────────────────────────────────────────────

def load_html(path: str) -> BeautifulSoup:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return BeautifulSoup(
            f.read(),
            "html.parser",
            store_line_numbers=True,
        )


def is_visible_text_node(node) -> bool:
    if not isinstance(node, NavigableString):
        return False

    if isinstance(node, (Comment, Doctype, Declaration, ProcessingInstruction)):
        return False

    if not str(node).strip():
        return False

    parent = node.parent

    while parent is not None and getattr(parent, "name", None):
        name = parent.name.lower()

        if name in SKIP_TEXT_PARENT_TAGS:
            return False

        if parent.get("aria-hidden") == "true":
            return False

        style = parent.get("style", "")
        style_norm = style.replace(" ", "").lower()

        if "display:none" in style_norm or "visibility:hidden" in style_norm:
            return False

        parent = parent.parent

    return True


def has_bold_ancestor(node) -> bool:
    parent = node.parent

    while parent is not None and getattr(parent, "name", None):
        name = parent.name.lower()

        if name in ("strong", "b", "h1", "h2", "h3", "h4", "h5", "h6"):
            return True

        style = parent.get("style", "")
        style_norm = style.replace(" ", "").lower()

        if (
            "font-weight:bold" in style_norm
            or "font-weight:700" in style_norm
            or "font-weight:800" in style_norm
            or "font-weight:900" in style_norm
        ):
            return True

        parent = parent.parent

    return False


def has_italic_ancestor(node) -> bool:
    parent = node.parent

    while parent is not None and getattr(parent, "name", None):
        name = parent.name.lower()

        if name in ("em", "i"):
            return True

        style = parent.get("style", "")
        style_norm = style.replace(" ", "").lower()

        if "font-style:italic" in style_norm:
            return True

        parent = parent.parent

    return False


def get_heading_level(node) -> int:
    parent = node.parent

    while parent is not None and getattr(parent, "name", None):
        name = parent.name.lower()

        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            return int(name[1])

        parent = parent.parent

    return 0


def is_inside_allowed_block(node) -> bool:
    parent = node.parent

    while parent is not None and getattr(parent, "name", None):
        if parent.name.lower() in ALLOWED_TEXT_BLOCKS:
            return True
        parent = parent.parent

    return False


def extract_html_tokens(soup: BeautifulSoup):
    tokens = []

    for node in soup.find_all(string=True):
        if not is_visible_text_node(node):
            continue

        text = strip_boilerplate_text(str(node))
        text = normalize_for_token_scan(text)
        if not text:
            continue

        line = getattr(node.parent, "sourceline", None)

        for m in TOKEN_RE.finditer(text):
            raw = m.group(0)

            if should_ignore_token(raw):
                continue

            tokens.append(
                TextToken(
                    text=raw,
                    norm=norm_token(raw),
                    line=line,
                    tag=node.parent,
                    bold=has_bold_ancestor(node),
                    italic=has_italic_ancestor(node),
                    heading_level=get_heading_level(node),
                    in_allowed_block=is_inside_allowed_block(node),
                )
            )

    return tokens


def _decode_data_uri_image(src: str):
    """Decode an inline data:image URI into a PIL image."""
    if not src or not src.lower().startswith("data:image/"):
        return None
    try:
        header, payload = src.split(",", 1)
        if ";base64" in header.lower():
            data = base64.b64decode(payload, validate=False)
        else:
            from urllib.parse import unquote_to_bytes
            data = unquote_to_bytes(payload)
        return Image.open(io.BytesIO(data)).convert("RGB")
    except Exception:
        return None


def _is_logo_placeholder_image(img) -> bool:
    """The Mytonomy HTML logo placeholder is a tiny 2x2 image."""
    try:
        return img is not None and img.width <= 8 and img.height <= 8
    except Exception:
        return False


def extract_html_images(soup: BeautifulSoup, html_path: str):
    """
    Extract HTML images from either:
      * data:image/...;base64,... inline images
      * normal relative image files
      * absolute local paths after server.py materialization.

    The Mytonomy logo placeholder is ignored later by check_images().
    """
    base = os.path.dirname(os.path.abspath(html_path))
    results = []

    for tag in soup.find_all("img"):
        src = tag.get("src", "")
        resolved = ""
        pil = None

        if src and src.lower().startswith("data:image/"):
            pil = _decode_data_uri_image(src)
            resolved = "data-uri"
        elif src and not src.startswith(("http://", "https://", "#")):
            resolved = os.path.normpath(os.path.join(base, src))
            if os.path.exists(resolved):
                try:
                    pil = Image.open(resolved).convert("RGB")
                except Exception:
                    pil = None

        results.append((tag, pil, resolved))

    return results


# ── PDF extraction ───────────────────────────────────────────────────────────

def _line_boxes_to_ignore(page):
    boxes = []

    for line in page.extract_text_lines() or []:
        text = _fix_rtl_line_order(line.get("text", ""))

        if is_boilerplate_line(text):
            boxes.append(
                (
                    float(line.get("top", 0)),
                    float(line.get("bottom", 0)),
                )
            )

    return boxes


def _is_inside_ignored_line_box(word, boxes) -> bool:
    top = float(word.get("top", 0))
    bottom = float(word.get("bottom", 0))

    for b_top, b_bottom in boxes:
        if top >= b_top - 2 and bottom <= b_bottom + 2:
            return True

    return False


def _is_rtl_text(text: str) -> bool:
    """True if the string is dominated by Arabic-script characters
    (Arabic/Farsi/Urdu). Used to detect PDFs where pdfplumber returns
    RTL glyphs in reversed visual order instead of logical reading order."""
    if not text:
        return False
    letters = re.findall(r"\w", text, flags=re.UNICODE)
    if not letters:
        return False
    rtl_count = len(ARABIC_SCRIPT_RE.findall(text))
    return rtl_count >= max(3, len(letters) * 0.4)


_RTL_RUN_SPLIT_RE = re.compile(
    r"([\u0600-\u065F\u066A-\u06EF\u06FA-\u06FF\u0750-\u077F\uFB50-\uFDFF\uFE70-\uFEFF]+"
    r"(?:\s+[\u0600-\u065F\u066A-\u06EF\u06FA-\u06FF\u0750-\u077F\uFB50-\uFDFF\uFE70-\uFEFF]+)*)"
)


def _fix_rtl_line_order(text: str) -> str:
    """pdfplumber returns Arabic/Farsi/Urdu glyphs character-reversed within
    each word/run, because these PDFs draw RTL text right-to-left but
    extraction reads left-to-right. Verified directly against the real
    source PDF: embedded Latin/digit tokens (©, 2025, Mytonomy, Inc., page
    numbers) are ALREADY correct and must be left untouched in their
    original position -- only the Arabic-script run(s) need their word
    order AND in-word character order reversed. A naive whole-line reversal
    is WRONG (it drags the Latin/digit tokens out of place)."""
    if not _is_rtl_text(text):
        return text

    runs = _RTL_RUN_SPLIT_RE.split(text)
    fixed_runs = []
    for run in runs:
        if not run:
            continue
        if _RTL_RUN_SPLIT_RE.fullmatch(run):
            fixed_runs.append(run[::-1])
        else:
            fixed_runs.append(run)
    return "".join(fixed_runs)


_RTL_WORD_CHAR_RE = re.compile(
    r"[\u0600-\u065F\u066A-\u06EF\u06FA-\u06FF\u0750-\u077F\uFB50-\uFDFF\uFE70-\uFEFF]"
)


def _is_rtl_word(text: str) -> bool:
    """True if the word contains an Arabic-script LETTER (not just an
    Arabic-Indic/Persian digit). Digits, Latin words, and punctuation-only
    tokens are already stored correctly by the PDF and must not be swept
    into the reversal, verified directly against real source PDFs."""
    return bool(_RTL_WORD_CHAR_RE.search(text or ""))


def _reorder_rtl_word_group(words, text_key="text"):
    """Given a list of word dicts (pdfplumber word objects) representing one
    visual line, reverse the order + in-word character order of each
    maximal run of consecutive Arabic-script words, leaving embedded
    Latin/digit/punctuation tokens (©, 2025, CGM, page numbers...) exactly
    where they are. Returns a new list of word dicts (shallow-copied where
    text changes) in corrected reading order."""
    out = []
    i = 0
    n = len(words)
    while i < n:
        if _is_rtl_word(str(words[i].get(text_key, ""))):
            j = i
            while j < n and _is_rtl_word(str(words[j].get(text_key, ""))):
                j += 1
            run = words[i:j]
            for w in reversed(run):
                w2 = dict(w)
                w2[text_key] = str(w.get(text_key, ""))[::-1]
                out.append(w2)
            i = j
        else:
            out.append(words[i])
            i += 1
    return out


def _build_pdf_words_from_chars(page, ignored_boxes):
    """Reconstruct words directly from character glyph positions instead of
    trusting pdfplumber's word-clustering/content-stream order, which is
    unreliable for RTL (Arabic/Farsi/Urdu) PDFs -- verified directly against
    real source files: extract_words(use_text_flow=True) shreds RTL runs
    into broken, mis-ordered fragments, while sorting characters by their
    actual x-position (right-to-left for RTL lines) reconstructs perfectly
    correct reading order, since glyph position -- not emission order -- is
    what's actually reliable in these PDFs."""
    chars = [c for c in page.chars if not _is_inside_ignored_line_box(c, ignored_boxes)]
    chars = [dict(c, _page_idx=idx) for idx, c in enumerate(chars)]

    # Combining marks (Arabic tashkeel etc.) sometimes render with a 'top'
    # value that rounds into a DIFFERENT line bucket than their base
    # letter -- verified directly against real source files. Build line
    # buckets from base (non-combining) characters first, then snap each
    # combining mark into whichever bucket is vertically closest, instead
    # of letting it form/join a bucket purely by its own rounded top.
    base_chars = [c for c in chars if unicodedata.combining((c.get("text") or "")[:1]) == 0]
    mark_chars = [c for c in chars if unicodedata.combining((c.get("text") or "")[:1]) != 0]

    line_groups = {}
    for c in base_chars:
        key = round(float(c.get("top", 0)), 0)
        line_groups.setdefault(key, []).append(c)

    bucket_keys = sorted(line_groups.keys())
    for c in mark_chars:
        raw_top = float(c.get("top", 0))
        if bucket_keys:
            key = min(bucket_keys, key=lambda k: abs(k - raw_top))
        else:
            key = round(raw_top, 0)
        line_groups.setdefault(key, []).append(c)

    # Restore true content-stream order within each bucket -- the mark
    # snapping above appended marks after all base chars, which would
    # otherwise break the adjacency the diacritic-anchoring logic below
    # depends on.
    for key in line_groups:
        line_groups[key] = sorted(line_groups[key], key=lambda c: c["_page_idx"])

    words = []
    all_lines = []  # list of line_words lists, one per line, built first so
                     # a single stable PAGE-level median gap can be computed
                     # (a per-line median is unreliable for short lines --
                     # e.g. a 2-word bullet item where the bullet glyph's own
                     # large isolation gap becomes the "median" of just 2
                     # samples -- verified directly against real source files)
    for key in sorted(line_groups):
        line_chars = line_groups[key]
        sample_text = "".join(c.get("text", "") for c in line_chars)
        is_rtl_line = _is_rtl_text(sample_text)

        # Tag original content-stream emission order BEFORE the x-position
        # sort below -- needed to correctly restore embedded Latin/digit
        # runs (see note further down).
        line_chars = [dict(c, _orig_idx=idx) for idx, c in enumerate(line_chars)]

        if is_rtl_line:
            # Combining marks (Arabic tashkeel: shadda, fatha, kasra,
            # damma, sukun, tanwin...) have unreliable OWN x-positions in
            # this PDF's export -- verified directly against real source
            # files. But their base letter is reliably identifiable: for
            # an RTL run, raw content-stream order is the exact reverse of
            # correct reading order (letters AND diacritics together), so
            # a mark's base letter is always the character immediately
            # AFTER it in raw content order. Anchor each mark's sort
            # position to sit just after (in descending-x order) that base
            # letter, instead of trusting the mark's own x0.
            by_orig_idx = {c["_orig_idx"]: c for c in line_chars}
            for c in line_chars:
                if unicodedata.combining((c.get("text") or "")[:1]) != 0:
                    base = by_orig_idx.get(c["_orig_idx"] + 1)
                    if base is not None:
                        c["x0"] = float(base.get("x0", c.get("x0", 0))) - 0.01

        line_chars = sorted(line_chars, key=lambda c: c.get("x0", 0), reverse=is_rtl_line)

        if is_rtl_line:
            # Sorting by x-position (right-to-left) correctly reconstructs
            # Arabic/Farsi letter order, but embedded Latin/digit runs
            # (acronyms like "SGLT2", "A1c", numbers like "150") are a
            # separate case: their OWN x-positions are unreliable in this
            # PDF's export (verified directly against real source files --
            # e.g. "SGLT2"'s digit '2' has a smaller x0 than 'S','G','L','T',
            # so sorting or even simple reversal both get it wrong), but
            # their ORIGINAL CONTENT-STREAM ORDER is correct. So: restore
            # each non-RTL run to its original emission order instead of
            # just reversing the x-sorted result.
            fixed = []
            i = 0
            n = len(line_chars)
            while i < n:
                ch_text = line_chars[i].get("text", "")
                if not _is_rtl_word(ch_text):
                    j = i
                    while j < n and not _is_rtl_word(line_chars[j].get("text", "")):
                        j += 1
                    run = sorted(line_chars[i:j], key=lambda c: c["_orig_idx"])
                    fixed.extend(run)
                    i = j
                else:
                    fixed.append(line_chars[i])
                    i += 1
            line_chars = fixed

        cur_text = []
        cur_size = None
        cur_font = None
        cur_start_edge = None  # position where this word's run started
        prev_edge = None  # trailing edge of the previous char, in reading direction
        line_words = []  # (text, size, fontname, start_edge, end_edge) for this line

        def flush():
            if cur_text:
                line_words.append({
                    "text": "".join(cur_text),
                    "top": key,
                    "size": cur_size,
                    "fontname": cur_font,
                    "start_edge": cur_start_edge,
                    "end_edge": prev_edge,
                })

        for c in line_chars:
            raw_char = c.get("text", "")
            if raw_char.strip() == "":
                # A literal space glyph in the PDF's character stream is,
                # in this document family, the most reliable word-boundary
                # signal available (position-gap alone is not reliable
                # enough to distinguish real word gaps from natural
                # non-connecting-letter gaps within a word) -- verified
                # directly against real source files. A small number of
                # PDFs embed a stray mid-word space (justification
                # artifact) which this will over-split; that residual is
                # far smaller than the alternative of losing real word
                # boundaries throughout the document.
                flush()
                cur_text = []
                cur_start_edge = None
                prev_edge = None
                continue

            size = float(c.get("size", 10) or 10)
            gap_thresh = max(1.5, size * 0.28)
            x0, x1 = float(c.get("x0", 0)), float(c.get("x1", 0))
            edge = x0 if is_rtl_line else x1

            # Unicode combining marks (Arabic tashkeel: shadda, fatha,
            # kasra, damma, sukun, tanwin...) are zero-width and decorate
            # the PRECEDING base letter -- verified directly against real
            # source files: their own bounding box often sits at a slightly
            # offset x-position that looks like a large physical gap to the
            # normal word-boundary check below, incorrectly splitting a
            # word right before its own diacritic (e.g. "كنتِ" -> "كن" +
            # "ِت"). They must never trigger a gap-based split.
            is_combining_mark = unicodedata.combining(raw_char[:1]) != 0 if raw_char else False

            gap = None
            if prev_edge is not None and not is_combining_mark:
                gap = (prev_edge - x1) if is_rtl_line else (x0 - prev_edge)

            if gap is not None and gap > gap_thresh:
                flush()
                cur_text = []
                cur_start_edge = None

            if cur_start_edge is None:
                cur_start_edge = x1 if is_rtl_line else x0
            cur_text.append(c.get("text", ""))
            cur_size = size
            cur_font = c.get("fontname", "")
            if not is_combining_mark:
                prev_edge = edge

        flush()
        all_lines.append(line_words)

    # Page-level median word-gap, from every line's word-boundary pairs,
    # excluding single-character words (bullet glyphs "•" etc.) whose own
    # isolation gap is not representative of normal inter-word spacing.
    page_gaps = []
    for line_words in all_lines:
        for i in range(1, len(line_words)):
            prev_w, cur_w = line_words[i - 1], line_words[i]
            if len(prev_w["text"]) <= 1 or len(cur_w["text"]) <= 1:
                continue
            if prev_w["end_edge"] is not None and cur_w["start_edge"] is not None:
                page_gaps.append(abs(cur_w["start_edge"] - prev_w["end_edge"]))
    page_gaps.sort()
    median_gap = page_gaps[len(page_gaps) // 2] if page_gaps else None

    # Post-merge pass, two related fixes for the same root cause (PDF
    # glyph-spacing quirks producing spurious extra word splits):
    #
    # 1. Some PDF exports embed a stray mid-word space glyph (RTL
    #    justification artifact) that has no reliable way to be told apart
    #    from a real space AT SPLIT TIME (verified directly against real
    #    source files -- ignoring the space's own bounding box entirely
    #    caused far worse mis-splits elsewhere in this same document). But
    #    AFTER splitting, a genuine phantom-space orphan is identifiable:
    #    it sits much closer to its neighbor than this page's typical word
    #    gap.
    # 2. Hyphenated compounds (English "mini-stroke", "X-rays",
    #    "blood-thinning"...) sometimes get extra kerning slack around the
    #    hyphen glyph itself, splitting into "mini" / "-" / "stroke" as
    #    three tokens -- verified directly against real source files. This
    #    is script-independent (pure Latin text), so it's handled
    #    separately from the RTL-word check below: a lone "-" token flanked
    #    by small gaps on both sides always re-fuses with its neighbors,
    #    regardless of script.
    for line_words in all_lines:
        if len(line_words) > 1 and median_gap:
            merged = [line_words[0]]
            for w in line_words[1:]:
                prev_w = merged[-1]
                gap = None
                if prev_w["end_edge"] is not None and w["start_edge"] is not None:
                    gap = abs(w["start_edge"] - prev_w["end_edge"])

                is_lone_hyphen = w["text"] in ("-", "\u2010", "\u2011")
                prev_ends_with_hyphen = prev_w["text"].endswith(("-", "\u2010", "\u2011"))
                looks_like_phantom_split = (
                    gap is not None and median_gap
                    and gap < median_gap * 0.3
                    and (_is_rtl_word(w["text"]) or is_lone_hyphen or prev_ends_with_hyphen)
                )
                if looks_like_phantom_split:
                    prev_w["text"] = prev_w["text"] + w["text"]
                    prev_w["end_edge"] = w["end_edge"]
                else:
                    merged.append(w)
            line_words = merged

        for w in line_words:
            words.append({
                "text": w["text"],
                "top": w["top"],
                "size": w["size"],
                "fontname": w["fontname"],
            })

    return words


def extract_pdf_text(pdf_path: str) -> str:
    parts = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            for line in page.extract_text_lines() or []:
                text = line.get("text", "")
                text = strip_boilerplate_text(text)
                text = normalize_for_token_scan(text)

                if not text:
                    continue

                text = _fix_rtl_line_order(text)
                parts.append(text)

    return "\n".join(parts)


def extract_pdf_tokens(pdf_path: str):
    tokens = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            ignored_boxes = _line_boxes_to_ignore(page)

            chars = page.chars or []
            usable_chars = [
                c for c in chars
                if not _is_inside_ignored_line_box(c, ignored_boxes)
            ]

            sizes = [round(float(c["size"])) for c in usable_chars if "size" in c]
            body_size = max(set(sizes), key=sizes.count) if sizes else 10

            try:
                words = page.extract_words(
                    keep_blank_chars=False,
                    use_text_flow=True,
                    extra_attrs=["fontname", "size"],
                )
            except Exception:
                words = page.extract_words(
                    keep_blank_chars=False,
                    use_text_flow=True,
                )

            # word-clustering / content-stream order from extract_words is
            # unreliable for RTL PDFs (verified against real source files),
            # so words are reconstructed from raw glyph positions instead.
            # This also makes the old line-grouping + reversal step
            # unnecessary -- reordered_words IS the corrected word list.
            reordered_words = _build_pdf_words_from_chars(page, ignored_boxes)

            for w in reordered_words:
                raw_word = strip_boilerplate_text(w.get("text", ""))
                raw_word = normalize_for_token_scan(raw_word)
                if not raw_word:
                    continue

                fontname = str(w.get("fontname", "")).lower()

                try:
                    size = float(w.get("size", body_size))
                except Exception:
                    size = body_size

                is_bold = (
                    "bold" in fontname
                    or "black" in fontname
                    or "heavy" in fontname
                    or "semibold" in fontname
                )

                heading_level = 0
                if size >= body_size * 1.55:
                    heading_level = 1
                elif size >= body_size * 1.30:
                    heading_level = 2
                elif size >= body_size * 1.15:
                    heading_level = 3

                for m in TOKEN_RE.finditer(raw_word):
                    raw = m.group(0)

                    if should_ignore_token(raw):
                        continue

                    tokens.append(
                        TextToken(
                            text=raw,
                            norm=norm_token(raw),
                            line=None,
                            tag=None,
                            bold=is_bold,
                            heading_level=heading_level,
                            in_allowed_block=True,
                        )
                    )

    return tokens


def extract_pdf_images(pdf_path: str):
    page_imgs = []

    with pdfplumber.open(pdf_path) as pdf:
        num_pages = len(pdf.pages)

        for page_idx, page in enumerate(pdf.pages):
            for obj in page.images:
                try:
                    width = obj["x1"] - obj["x0"]
                    height = obj["y1"] - obj["y0"]
                    top = page.height - obj["y1"]

                    # Ignore header logo / blank logo placeholder area.
                    if (
                        obj["x0"] <= page.width * HEADER_LOGO_MAX_LEFT_RATIO
                        and top <= page.height * HEADER_LOGO_MAX_TOP_RATIO
                        and width <= page.width * HEADER_LOGO_MAX_WIDTH_RATIO
                        and height <= page.height * HEADER_LOGO_MAX_HEIGHT_RATIO
                    ):
                        continue

                    bbox = (
                        obj["x0"],
                        page.height - obj["y1"],
                        obj["x1"],
                        page.height - obj["y0"],
                    )

                    cropped = page.crop(bbox).to_image(resolution=100).original
                    cropped = cropped.convert("RGB")

                    if is_blank_or_logo_placeholder(cropped):
                       continue

# Store original PDF display dimensions.
# These are the width/height of the image placement inside the PDF page.
                    cropped._qa_pdf_display_width = float(width)
                    cropped._qa_pdf_display_height = float(height)

                    page_imgs.append((page_idx, cropped))

                    

                except Exception:
                    continue

    if not page_imgs:
        return []

    repeat_threshold = max(2, math.ceil(num_pages / 2))
    groups = []

    for page_idx, img in page_imgs:
        try:
            h = imagehash.phash(img)
        except Exception:
            continue

        for group in groups:
            if (h - group[0]) <= 8:
                group[1].add(page_idx)
                break
        else:
            groups.append((h, {page_idx}, img))

    return [img for _, pages, img in groups if len(pages) < repeat_threshold]
    
def parse_css_length(value):
    if value is None:
        return None

    value = str(value).strip().lower()

    m = re.search(r"([0-9]+(?:\.[0-9]+)?)", value)
    if not m:
        return None

    try:
        return float(m.group(1))
    except Exception:
        return None


def get_img_declared_dimensions(tag):
    style = parse_inline_style(tag.get("style", ""))

    width = parse_css_length(tag.get("width"))
    height = parse_css_length(tag.get("height"))

    if width is None:
        width = parse_css_length(style.get("width"))

    if height is None:
        height = parse_css_length(style.get("height"))

    return width, height


def _ratio_error(actual, expected):
    if not actual or not expected:
        return 0.0

    return abs(actual - expected) / expected


def check_image_dimensions(tag, matched_source_img, html_img, line):
    issues = []

    declared_w, declared_h = get_img_declared_dimensions(tag)

    # No declared width/height means there is nothing to validate.
    if declared_w is None and declared_h is None:
        return issues

    # Some extractors carry image metadata alongside the image itself, e.g.
    # (page_number, image). Coerce defensively so a metadata-wrapped value
    # here can never blow up the whole file's QA — matches _coerce_pil_image
    # used elsewhere for the same reason.
    source_meta = matched_source_img
    matched_source_img = _coerce_pil_image(matched_source_img)
    html_img = _coerce_pil_image(html_img)
    if matched_source_img is None or html_img is None:
        return issues

    expected_w = getattr(source_meta, "_qa_source_display_width", None) or getattr(matched_source_img, "_qa_source_display_width", None)
    expected_h = getattr(source_meta, "_qa_source_display_height", None) or getattr(matched_source_img, "_qa_source_display_height", None)
    if not expected_w or not expected_h:
        expected_w = getattr(source_meta, "_qa_pdf_display_width", None) or getattr(matched_source_img, "_qa_pdf_display_width", None)
        expected_h = getattr(source_meta, "_qa_pdf_display_height", None) or getattr(matched_source_img, "_qa_pdf_display_height", None)

    # Fallback to image pixel ratio if PDF display size is unavailable.
    if not expected_w or not expected_h:
        expected_w = matched_source_img.width
        expected_h = matched_source_img.height

    if not expected_h:
        return issues

    expected_ratio = expected_w / expected_h

    # 8% tolerance for PDF-point vs HTML-pixel conversion differences.
    size_tolerance = 0.08

    size_mismatch = False
    mismatch_parts = []

    if declared_w is not None:
        w_error = _ratio_error(declared_w, expected_w)

        if w_error > size_tolerance:
            size_mismatch = True
            mismatch_parts.append(
                f"width expected about {expected_w:.1f}, got {declared_w:.1f}"
            )

    if declared_h is not None:
        h_error = _ratio_error(declared_h, expected_h)

        if h_error > size_tolerance:
            size_mismatch = True
            mismatch_parts.append(
                f"height expected about {expected_h:.1f}, got {declared_h:.1f}"
            )

    if size_mismatch:
        issues.append(Issue(
            "Image Display Size Mismatch",
            "error",
            line,
            "HTML image declared width/height differs from the source document image placement.",
            snippet=str(tag)[:300],
            expected=f"{expected_w:.1f} x {expected_h:.1f}",
            actual=(
                f"{declared_w if declared_w is not None else 'not set'} x "
                f"{declared_h if declared_h is not None else 'not set'}"
            ),
        ))

    # If both width and height are declared, also check distortion/aspect ratio.
    if declared_w is not None and declared_h is not None and declared_h != 0:
        declared_ratio = declared_w / declared_h
        ratio_error = _ratio_error(declared_ratio, expected_ratio)

        if ratio_error > 0.03:
            issues.append(Issue(
                "Image Aspect Ratio Mismatch",
                "error",
                line,
                "HTML image width/height changes the expected image aspect ratio.",
                snippet=str(tag)[:300],
                expected=f"aspect ratio about {expected_ratio:.3f}",
                actual=f"{declared_w:.1f} x {declared_h:.1f}, ratio {declared_ratio:.3f}",
            ))

    # Check actual image file ratio too.
    if html_img and html_img.height:
        actual_file_ratio = html_img.width / html_img.height
        file_ratio_error = _ratio_error(actual_file_ratio, expected_ratio)

        if file_ratio_error > 0.03:
            issues.append(Issue(
                "Image File Aspect Ratio Mismatch",
                "error",
                line,
                "HTML image file aspect ratio differs from the matched source image.",
                snippet=str(tag)[:300],
                expected=f"aspect ratio about {expected_ratio:.3f}",
                actual=f"{html_img.width} x {html_img.height}, ratio {actual_file_ratio:.3f}",
            ))

    return issues
    
def check_title(pdf_path: str, soup: BeautifulSoup):
    issues = []

    expected_title = extract_pdf_title(pdf_path)
    title_tag = soup.find("title")

    if not expected_title:
        return issues

    if not title_tag:
        issues.append(Issue(
            "Missing Title Tag",
            "warning",
            None,
            "HTML is missing a <title> tag.",
            expected=expected_title,
            actual="",
        ))
        return issues

    actual_title = title_tag.get_text(" ", strip=True)

    if norm_text(expected_title) != norm_text(actual_title):
        issues.append(Issue(
            "Title Mismatch",
            "error",
            getattr(title_tag, "sourceline", None),
            f'Expected HTML title "{expected_title}" but found "{actual_title}".',
            expected=expected_title,
            actual=actual_title,
            snippet=f"<title>{actual_title}</title>",
        ))

    return issues


def _coerce_pil_image(img):
    """Return a PIL Image from an image value used by the QA engine.

    Some extractors store image metadata together with the image, e.g.
    ``(page_number, image)``.  The comparison layer should accept either
    representation instead of trying to call ``convert()`` on the tuple.
    """
    if isinstance(img, Image.Image):
        return img

    if isinstance(img, (tuple, list)):
        # Common forms: (page_no, PIL.Image) or (metadata, PIL.Image).
        for value in reversed(img):
            if isinstance(value, Image.Image):
                return value

    # Some PIL-compatible objects expose convert() without being an Image.Image.
    if hasattr(img, "convert"):
        return img

    return None


def _hash_image(img):
    """Hash an image safely, including extractor values carrying metadata."""
    pil = _coerce_pil_image(img)
    if pil is None:
        return None
    pil = pil.convert("RGB")
    return imagehash.phash(pil), imagehash.colorhash(pil)


def check_images(source_imgs, html_img_entries, hash_threshold=18, color_threshold=15):
    """
    Compare source-document images against HTML images.

    The comparison is document-type agnostic: PDF and RTF source images are
    represented as PIL images with optional source display dimensions.
    """
    issues = []

    valid_html = []

    for tag, pil, resolved in html_img_entries:
        src = tag.get("src", "(no src)")
        basename = os.path.basename(src).lower()

        # Ignore Mytonomy's recurring logo, including its inline 2x2 data URI.
        if basename in IGNORE_IMAGES or _is_logo_placeholder_image(pil):
            continue

        if pil is None:
            line = getattr(tag, "sourceline", None)
            issues.append(
                Issue(
                    "Broken Image",
                    "error",
                    line,
                    f"Image file/data URI could not be decoded: {src[:180]}",
                    expected="Readable image",
                    actual=src[:180],
                )
            )
            continue

        valid_html.append((tag, pil, resolved))

    # No source images means there is no source image inventory to compare.
    if not source_imgs:
        return issues

    valid_source = []
    for img in source_imgs:
        if img is None:
            continue
        if _is_logo_placeholder_image(img):
            continue
        if getattr(img, "_qa_source_name", "").lower() in IGNORE_IMAGES:
            continue
        valid_source.append(img)

    if not valid_source:
        return issues

    src_hashes = [_hash_image(img) for img in valid_source]
    src_hashes = [h for h in src_hashes if h is not None]
    html_hashes = [_hash_image(pil) for _, pil, _ in valid_html]
    html_hashes = [h for h in html_hashes if h is not None]

    if not src_hashes:
        return issues

    src_phashes = [h[0] for h in src_hashes]
    src_colorhashes = [h[1] for h in src_hashes]

    matched_src = set()
    order_seq = []

    for idx_html, ((tag, pil, resolved), hash_pair) in enumerate(zip(valid_html, html_hashes)):
        h_p, h_c = hash_pair

        p_dists = [h_p - sh for sh in src_phashes]
        best_i = int(np.argmin(p_dists))
        best_p_dist = p_dists[best_i]
        c_dist = h_c - src_colorhashes[best_i]
        line = getattr(tag, "sourceline", None)

        if best_p_dist <= hash_threshold and c_dist <= color_threshold:
            matched_src.add(best_i)
            order_seq.append((best_i, tag, line))

            issues += check_image_dimensions(
                tag=tag,
                matched_source_img=valid_source[best_i],
                html_img=pil,
                line=line,
            )

        elif best_p_dist <= hash_threshold:
            issues.append(
                Issue(
                    "Image Possibly Wrong",
                    "warning",
                    line,
                    "Image structure matches a source image, but colors differ. Please verify manually.",
                    snippet=resolved,
                    expected="source image color profile",
                    actual="HTML image color profile differs",
                )
            )
            matched_src.add(best_i)
            order_seq.append((best_i, tag, line))

            issues += check_image_dimensions(
                tag=tag,
                matched_source_img=valid_source[best_i],
                html_img=pil,
                line=line,
            )

        else:
            similarity = max(0.0, 1.0 - best_p_dist / 64.0)
            issues.append(
                Issue(
                    "Image Doesn't Match Source",
                    "error",
                    line,
                    f"This image does not match any image in the source. Best visual similarity: {similarity:.0%}.",
                    snippet=resolved,
                    expected="Matching source image",
                    actual=os.path.basename(resolved) if resolved != "data-uri" else "inline data image",
                )
            )

    for idx in range(len(valid_source)):
        if idx not in matched_src:
            issues.append(
                Issue(
                    "Missing Image",
                    "error",
                    None,
                    f"Source image #{idx + 1} was not found in the HTML.",
                    expected=f"source image #{idx + 1}",
                    actual="Missing in HTML",
                )
            )

    # Image Order Mismatch: images that all matched correctly by content
    # but appear in a different sequence than the source document. Only
    # meaningful when every matched image is unique (a repeated image can't
    # have a single "correct" position) and there are at least 2 to order.
    matched_src_indices = [pair[0] for pair in order_seq]
    if len(matched_src_indices) >= 2 and len(set(matched_src_indices)) == len(matched_src_indices):
        if matched_src_indices != sorted(matched_src_indices):
            first_bad_pos = next(
                i for i in range(len(matched_src_indices))
                if matched_src_indices[i] != sorted(matched_src_indices)[i]
            )
            _, bad_tag, bad_line = order_seq[first_bad_pos]
            issues.append(
                Issue(
                    "Image Order Mismatch",
                    "error",
                    bad_line,
                    "Images match the source document's images, but appear in a different "
                    f"order. Expected source order {sorted(matched_src_indices)}, "
                    f"got {matched_src_indices}.",
                    snippet=str(bad_tag)[:300],
                    expected=f"source image order {sorted(matched_src_indices)}",
                    actual=f"HTML image order {matched_src_indices}",
                )
            )

    # Extra HTML images are images which were not mapped to a source image.
    # Do not report an HTML image twice if it already generated a mismatch.
    matched_html_count = len(order_seq)
    if len(valid_html) > len(valid_source):
        for idx in range(len(valid_source), len(valid_html)):
            tag, _, resolved = valid_html[idx]
            line = getattr(tag, "sourceline", None)
            issues.append(
                Issue(
                    "Extra Image",
                    "error",
                    line,
                    "HTML contains an image that has no corresponding source image.",
                    expected="No extra image",
                    actual=resolved if resolved else "inline data image",
                )
            )

    return issues

def has_b_or_strong_descendant(tag) -> bool:
    if not tag:
        return False

    # If the tag itself is b/strong
    if getattr(tag, "name", "").lower() in ("b", "strong"):
        return True

    return tag.find(["b", "strong"]) is not None


def text_prefix_before_colon(text: str) -> str:
    """
    Returns label prefix like:
      'Pituitary tumor:'
      'Blood tests:'
    """
    text = WS.sub(" ", text or "").strip()

    m = re.match(r"^(.{1,90}?:)", text)
    if not m:
        return ""

    return m.group(1).strip()
    
def check_required_b_tags_by_company_spec(soup: BeautifulSoup):
    """
    Company-structure rule:
    Only label-style text inside list paragraphs should require <b>/<strong>.

    Example:
      <p><b>Blood tests:</b> remaining text</p>
    """
    issues = []
    seen = set()

    def add_issue(tag, message, expected="", actual=""):
        line = getattr(tag, "sourceline", None)
        snippet = str(tag)[:350]
        key = ("Missing B Tag", line, snippet, expected, actual)

        if key in seen:
            return

        seen.add(key)
        issues.append(Issue(
            "Missing B Tag",
            "error",
            line,
            message,
            snippet=snippet,
            expected=expected,
            actual=actual,
        ))

    for li in soup.find_all("li"):
        for p in li.find_all("p", recursive=True):
            text = p.get_text(" ", strip=True)
            label = text_prefix_before_colon(text)

            if not label:
                continue

            label_is_bold = False

            for btag in p.find_all(["b", "strong"]):
                bold_text = btag.get_text(" ", strip=True)

                if label.lower() == bold_text.lower():
                    label_is_bold = True
                    break

            if not label_is_bold:
                add_issue(
                    p,
                    f'List label "{label}" should be wrapped in <b>/<strong>.',
                    expected=f"<b>{label}</b>",
                    actual=label,
                )

    return issues

def check_missing_b_tags_strict(soup: BeautifulSoup):
    """
    Strict company HTML structure check for removed <b>/<strong> tags.

    Only checks h1/h2/h3 headings. It does NOT check every colon label,
    because labels like "Benefits:" and "Call your care team if:" are valid
    converter output in your current files and should not create noise.

    Catches:
      <h2><b>Before the Procedure</b></h2>
    changed to:
      <h2>Before the Procedure</h2>
    """
    issues = []
    seen = set()

    def add_issue(tag, message, expected, actual):
        line = getattr(tag, "sourceline", None)
        snippet = str(tag)[:350]
        key = ("Missing B Tag", line, snippet, expected, actual)

        if key in seen:
            return

        seen.add(key)
        issues.append(Issue(
            "Missing B Tag",
            "error",
            line,
            message,
            snippet=snippet,
            expected=expected,
            actual=actual,
        ))

    # h1-h3 headings are expected to contain explicit <b>/<strong> in this converter output.
    # h4-h6 are skipped because real clean files may use them without inner <b>.
    for heading in soup.find_all(["h1", "h2", "h3"]):
        text = heading.get_text(" ", strip=True)
        if not text:
            continue

        if heading.find(["b", "strong"]) is None:
            add_issue(
                heading,
                f"<{heading.name}> heading text is missing the required <b>/<strong> wrapper.",
                f"<{heading.name}><b>{text}</b></{heading.name}>",
                str(heading)[:200],
            )

    return issues

def check_tag_attributes_and_styles(soup: BeautifulSoup):
    issues = []
    seen = set()

    def add_issue(category, tag, message, expected="", actual=""):
        line = getattr(tag, "sourceline", None)
        snippet = str(tag)[:350]

        key = (category, line, snippet, expected, actual)

        if key in seen:
            return

        seen.add(key)

        issues.append(Issue(
            category,
            "error",
            line,
            message,
            snippet=snippet,
            expected=expected,
            actual=actual,
        ))

    # Check <li data-list-text="">
    for li in soup.find_all("li"):
        text = li.get_text(" ", strip=True)

        if not text:
            continue

        marker = (li.get("data-list-text") or "").strip()

        if marker != EXPECTED_BULLET_MARKER:
            add_issue(
                "Tag Attribute Mismatch",
                li,
                "List item data-list-text attribute differs from the expected bullet marker.",
                expected=f'data-list-text="{EXPECTED_BULLET_MARKER}"',
                actual=f'data-list-text="{marker}"',
            )

    # Check <p> styles inside list items
    for li in soup.find_all("li"):
        for p in li.find_all("p", recursive=True):
            text = p.get_text(" ", strip=True)

            if not text:
                continue

            style = parse_inline_style(p.get("style", ""))

            expected_styles = {
                "padding-left": "23pt",
                "text-indent": "-10pt",
                "line-height": "110%",
                "text-align": "left",
            }

            for prop, expected_value in expected_styles.items():
                actual_value = style.get(prop)

                if actual_value != expected_value:
                    add_issue(
                        "Style Attribute Mismatch",
                        p,
                        f'List paragraph style "{prop}" differs from expected value.',
                        expected=f"{prop}: {expected_value}",
                        actual=f"{prop}: {actual_value if actual_value is not None else 'missing'}",
                    )

            if "padding-right" in style:
                add_issue(
                    "Style Attribute Mismatch",
                    p,
                    "List paragraph uses padding-right, but expected padding-left.",
                    expected="padding-left: 23pt",
                    actual=f"padding-right: {style.get('padding-right')}",
                )

    return issues
    
def is_hyphen_token_text(text: str) -> bool:
    return (text or "").strip() in HYPHEN_TOKENS


def merge_spaced_hyphen_tokens(tokens):
    """
    Converts:
      long - term
      half - gallon

    into:
      long-term
      half-gallon

    This avoids false Changed Text errors when PDF extraction adds spaces
    around hyphens but HTML has normal hyphenated words.
    """

    def is_word_token(text):
        return bool(uregex.fullmatch(r"[\p{L}\p{M}\p{N}]+(?:\'[\p{L}\p{M}\p{N}]+)?", text or ""))

    def is_hyphen_token(text):
        return (text or "").strip() in HYPHEN_TOKENS

    merged = []
    i = 0

    while i < len(tokens):
        if (
            i + 2 < len(tokens)
            and is_word_token(tokens[i].text)
            and is_hyphen_token(tokens[i + 1].text)
            and is_word_token(tokens[i + 2].text)
        ):
            combined_text = f"{tokens[i].text}-{tokens[i + 2].text}"

            merged.append(TextToken(
                text=combined_text,
                norm=norm_token(combined_text),
                line=tokens[i].line,
                tag=tokens[i].tag,
                bold=tokens[i].bold or tokens[i + 2].bold,
                italic=tokens[i].italic or tokens[i + 2].italic,
                heading_level=tokens[i].heading_level or tokens[i + 2].heading_level,
                in_allowed_block=(
                    tokens[i].in_allowed_block
                    and tokens[i + 2].in_allowed_block
                ),
            ))

            i += 3
            continue

        merged.append(tokens[i])
        i += 1

    return merged

# ── Strict text comparison ───────────────────────────────────────────────────

def _qa_hyphen_cmp(text: str) -> str:
    """Normalize only hyphen spacing noise for final issue filtering."""
    text = normalize_for_compare(text)
    text = uregex.sub(r"[-‐‒–—−]", "-", text)
    text = uregex.sub(r"(?<=[\p{L}\p{M}\p{N}])\s*-\s*(?=[\p{L}\p{M}\p{N}])", "-", text)
    text = uregex.sub(r"\s+", " ", text).strip()
    return text


def _is_hyphen_spacing_noise(expected: str, actual: str) -> bool:
    if not expected or not actual:
        return False

    e = _qa_hyphen_cmp(expected)
    a = _qa_hyphen_cmp(actual)

    return e == a and e != str(expected or "").lower().strip() and "-" in e




def _qa_hyphen_canonical_for_drop(text: str) -> str:
    """
    Canonical form used only for suppressing false Changed Text issues caused by
    PDF extraction spaces around hyphens.
    """
    text = normalize_for_compare(text)
    text = uregex.sub(r"[-‐‒–—−]", "-", text)
    text = uregex.sub(r"(?<=[\p{L}\p{M}\p{N}])\s*-\s*(?=[\p{L}\p{M}\p{N}])", "-", text)
    text = uregex.sub(r"\s+", " ", text).strip().rstrip(".")
    return text


def _should_ignore_hyphen_only_issue(issue: Issue) -> bool:
    expected = str(getattr(issue, "expected", "") or "").strip()
    actual = str(getattr(issue, "actual", "") or "").strip()

    if not expected or not actual:
        return False

    expected_norm = _qa_hyphen_canonical_for_drop(expected)
    actual_norm = _qa_hyphen_canonical_for_drop(actual)

    if expected_norm != actual_norm:
        return False

    # Only suppress if there really was a hyphen spacing/dash representation difference.
    raw_pair = (expected + " " + actual).lower()
    return "-" in expected_norm and (
        " - " in raw_pair
        or "- " in raw_pair
        or " -" in raw_pair
        or "–" in raw_pair
        or "—" in raw_pair
        or "−" in raw_pair
        or "‐" in raw_pair
        or "‒" in raw_pair
    )

def _add_issue(issues, seen, issue: Issue):
    # Global guard for PDF extraction / converter hyphen spacing noise.
    # Examples: mini - stroke vs mini-stroke, X - rays vs X-rays.
    if (
        _should_ignore_hyphen_only_issue(issue)
        or _is_hyphen_spacing_noise(issue.expected, issue.actual)
        or is_hyphen_spacing_only_difference(issue.expected, issue.actual)
    ):
        return

    key = (
        issue.category,
        issue.line,
        issue.message,
        issue.expected,
        issue.actual,
    )

    if key not in seen:
        seen.add(key)
        issues.append(issue)

def extract_pdf_title(pdf_path: str) -> str:
    """
    Gets the first real non-footer PDF line.
    For this document type, that is usually the document title.
    """
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            lines = page.extract_text_lines() or []

            for line in lines:
                text = strip_boilerplate_text(line.get("text", ""))

                if not text:
                    continue

                return WS.sub(" ", text).strip()

    return ""

def is_hyphen_spacing_only_difference(expected: str, actual: str) -> bool:
    if not expected or not actual:
        return False

    return canonical_hyphen_text(expected) == canonical_hyphen_text(actual)


# ── Source document extraction helpers ───────────────────────────────────────


def _rtf_balanced_group(raw: str, start: int):
    """Return (group_text, end_index) for a balanced RTF group."""
    depth = 0
    i = start
    n = len(raw)
    while i < n:
        c = raw[i]
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                return raw[start:i + 1], i + 1
        elif c == "\\":
            # A backslash does not escape braces in RTF; braces are still structural.
            pass
        i += 1
    return raw[start:], n


def _strip_rtf_noncontent_groups(raw: str) -> str:
    """
    Remove RTF destinations that are not document content.

    In particular, INCLUDEPICTURE fields and \\pict groups must never become
    text tokens. The supplied Mytonomy RTFs use external INCLUDEPICTURE fields,
    with the actual image files in the sibling Images directory.
    """
    destinations = {
        "fonttbl", "stylesheet", "info", "colortbl", "filetbl",
        "listtable", "listoverridetable", "themedata", "xmlnstbl",
    }

    out = []
    i = 0
    n = len(raw)

    while i < n:
        if raw[i] == "{":
            group, nxt = _rtf_balanced_group(raw, i)
            prefix = group[:160]

            # The RTF root group ({\rtf1 ...} spanning the whole document)
            # must never be treated as a strippable destination. Without
            # this guard, _rtf_balanced_group(raw, 0) returns the entire
            # file as "group" (the outermost brace only closes at EOF), and
            # if \field/INCLUDEPICTURE/\pict happens to appear anywhere in
            # that whole-document string, the destination checks below would
            # match on the *root* group and silently strip the entire
            # document to an empty string. Only genuinely nested destination
            # groups (fonttbl, a specific \field, a specific \pict, etc.)
            # should ever be matched and skipped here.
            if re.match(r"\{\s*\\rtf\d*\b", prefix, re.I):
                out.append(raw[i])
                i += 1
                continue

            # Any field containing INCLUDEPICTURE is an image placeholder,
            # not document text.
            if re.search(r"\\field\b", prefix, re.I) and re.search(
                r"INCLUDEPICTURE", group, re.I
            ):
                i = nxt
                continue

            # Embedded image destination.
            if re.search(r"\\pict\b", prefix, re.I):
                i = nxt
                continue

            # Non-content metadata destinations.
            m = re.match(r"\{\s*(?:\\\*)?\\([A-Za-z]+)", prefix)
            if m and m.group(1).lower() in destinations:
                i = nxt
                continue

        out.append(raw[i])
        i += 1

    return "".join(out)


def _rtf_codepage(raw: str, rtf_path: str = "") -> str:
    # These Mytonomy RTFs declare ansicpg1252 but store Arabic/Persian/Urdu
    # fallback bytes using the Windows-1256 character mapping.
    rtl_lang_ids = {
        "1025", "2049", "3073", "4097", "5121", "6145", "7169",
        "8193", "9217", "10241", "11265", "12289", "13313", "14337",
        "15361", "16385", "1056", "1065",
    }
    if re.search(r"\\lang(?:fe)?(?:%s)\b" % "|".join(rtl_lang_ids), raw, re.I):
        return "cp1256"

    if rtf_path and re.search(r"(?:_| )(?:AR|FA|UR)\.rtf$", str(rtf_path), re.I):
        return "cp1256"

    m = re.search(r"\\ansicpg(\d+)", raw, re.I)
    if not m:
        return "cp1252"

    cp = m.group(1)
    try:
        "".encode(f"cp{cp}")
        return f"cp{cp}"
    except Exception:
        return "cp1252"


def _basic_rtf_to_text(raw: str, rtf_path: str = "") -> str:
    """
    Dependency-free RTF text parser used when striprtf is unavailable.

    Handles:
      * \\uN Unicode escapes + \\ucN fallback bytes
      * RTF hex bytes
      * paragraph/line/tab controls
      * common punctuation controls
      * removal of image and metadata groups

    Literal CR/LF characters in an RTF file are formatting whitespace, not
    document newlines, so they are discarded.
    """
    raw = _strip_rtf_noncontent_groups(raw)
    codepage = _rtf_codepage(raw, rtf_path)
    uc = 1
    out = []
    i = 0
    n = len(raw)

    control_symbols = {
        "~": "\u00a0",
        "-": "\u00ad",
        "_": "\u2011",
    }
    control_words = {
        "par": "\n",
        "line": "\n",
        "tab": "\t",
        "cell": "\t",
        "row": "\n",
        "emdash": "—",
        "endash": "–",
        "bullet": "•",
        "lquote": "‘",
        "rquote": "’",
        "ldblquote": "“",
        "rdblquote": "”",
    }

    while i < n:
        c = raw[i]

        if c in "\r\n":
            i += 1
            continue

        if c in "{}":
            i += 1
            continue

        if c != "\\":
            out.append(c)
            i += 1
            continue

        if i + 1 >= n:
            break

        nxt = raw[i + 1]

        # Escaped literal character.
        if nxt in "{}\\":
            out.append(nxt)
            i += 2
            continue

        # Hex byte: \\'hh
        if nxt == "'":
            hx = raw[i + 2:i + 4]
            if len(hx) == 2:
                try:
                    out.append(bytes.fromhex(hx).decode(codepage, errors="replace"))
                except Exception:
                    pass
            i += 4
            continue

        # Control symbol.
        if not nxt.isalpha():
            out.append(control_symbols.get(nxt, ""))
            i += 2
            continue

        # Control word.
        j = i + 1
        while j < n and raw[j].isalpha():
            j += 1

        word = raw[i + 1:j]
        sign = 1

        if j < n and raw[j] in "+-":
            if raw[j] == "-":
                sign = -1
            j += 1

        k = j
        while k < n and raw[k].isdigit():
            k += 1

        number = int(raw[j:k]) * sign if k > j else None

        # The delimiter space belongs to the control word.
        if k < n and raw[k] == " ":
            k += 1

        if word.lower() == "uc" and number is not None:
            uc = max(0, number)
            i = k
            continue

        if word.lower() == "u" and number is not None:
            value = number + 65536 if number < 0 else number
            out.append(chr(value))

            # Skip the fallback representation specified by \\ucN.
            q = k
            skipped = 0

            while skipped < uc and q < n:
                if raw[q] in "\r\n":
                    q += 1
                    continue

                if raw[q] == "\\" and q + 1 < n and raw[q + 1] == "'":
                    q += 4
                elif raw[q] == "\\":
                    qq = q + 1
                    while qq < n and raw[qq].isalpha():
                        qq += 1
                    while qq < n and raw[qq].isdigit():
                        qq += 1
                    if qq < n and raw[qq] == " ":
                        qq += 1
                    q = qq
                else:
                    q += 1

                skipped += 1

            i = q
            continue

        mapped = control_words.get(word.lower())
        if mapped is not None:
            out.append(mapped)

        i = k

    return "".join(out)


def _rtf_needs_codepage_override(raw: str, rtf_path: str = "") -> bool:
    """
    True when this RTF matches one of the conditions where _rtf_codepage()
    overrides the declared \\ansicpg with cp1256 — i.e. an RTL \\lang ID is
    present, or the filename carries an AR/FA/UR suffix. striprtf has no
    knowledge of this override, so these documents must go through the
    codepage-aware fallback parser instead, or their non-Unicode-aware
    fallback bytes (Arabic/Persian/Urdu) come out as mojibake.
    """
    rtl_lang_ids = {
        "1025", "2049", "3073", "4097", "5121", "6145", "7169",
        "8193", "9217", "10241", "11265", "12289", "13313", "14337",
        "15361", "16385", "1056", "1065",
    }
    if re.search(r"\\lang(?:fe)?(?:%s)\b" % "|".join(rtl_lang_ids), raw, re.I):
        return True
    if rtf_path and re.search(r"(?:_| )(?:AR|FA|UR)\.rtf$", str(rtf_path), re.I):
        return True
    return False


def extract_rtf_text(rtf_path: str) -> str:
    """
    Extract visible RTF document text.

    striprtf is used as the default parser for documents that don't need a
    codepage override. However, striprtf has no awareness of the
    ansicpg-vs-actual-encoding mismatch that Mytonomy's Arabic/Persian/Urdu
    RTFs exhibit (declared ansicpg1252 but hex-fallback bytes are cp1256) —
    it decodes those bytes with the wrong codepage and produces mojibake.
    For any document flagged by _rtf_needs_codepage_override(), the
    dependency-free, codepage-aware parser is used instead so this class of
    document is decoded correctly. Image fields are stripped before parsing
    either way so INCLUDEPICTURE/pict payloads never contaminate the text.
    """
    raw_bytes = Path(rtf_path).read_bytes()
    raw = raw_bytes.decode("latin1", errors="replace")
    content_raw = _strip_rtf_noncontent_groups(raw)

    if _rtf_needs_codepage_override(raw, rtf_path):
        text = _basic_rtf_to_text(raw, rtf_path)
    else:
        try:
            from striprtf.striprtf import rtf_to_text
            text = rtf_to_text(content_raw)
        except Exception:
            text = _basic_rtf_to_text(raw, rtf_path)

    parts = []
    for line in text.splitlines():
        line = strip_boilerplate_text(line)
        line = normalize_for_token_scan(line)
        if line:
            parts.append(line)

    return "\n".join(parts)


def _rtf_image_references(rtf_path: str):
    """Return external INCLUDEPICTURE filenames in RTF document order."""
    raw = Path(rtf_path).read_bytes().decode("latin1", errors="replace")
    refs = []

    for m in re.finditer(
        r"INCLUDEPICTURE\s+(?:\"([^\"]+)\"|([^\s\\}]+))",
        raw,
        re.I,
    ):
        ref = m.group(1) or m.group(2) or ""
        if ref:
            refs.append(ref.replace("\\", "/"))

    return refs


def _find_rtf_asset(rtf_path: str, ref: str):
    """Resolve an RTF INCLUDEPICTURE reference in the temp package."""
    source_dir = Path(rtf_path).resolve().parent
    basename = Path(ref).name

    candidates = [
        source_dir / ref,
        source_dir / "Images" / basename,
    ]

    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate

    # Server materialization may flatten the source while retaining Images/.
    matches = list(source_dir.rglob(basename))
    return matches[0] if matches else None


def _extract_embedded_rtf_images(rtf_path: str):
    """
    Best-effort extraction for truly embedded RTF \\pict blocks.
    External INCLUDEPICTURE assets are handled separately.
    """
    raw = Path(rtf_path).read_bytes().decode("latin1", errors="replace")
    images = []

    for m in re.finditer(r"\\pict\b", raw, re.I):
        # Walk back to the opening brace of this pict group.
        start = raw.rfind("{", 0, m.start())
        if start < 0:
            continue

        group, _ = _rtf_balanced_group(raw, start)
        if "\\pict" not in group.lower():
            continue

        png = bool(re.search(r"\\pngblip\b", group, re.I))
        jpeg = bool(re.search(r"\\jpe?gblip\b", group, re.I))
        if not (png or jpeg):
            continue

        hex_data = re.sub(r"[^0-9A-Fa-f]", "", group[m.end() - start:])
        # The group metadata contains hex-like numbers in control words, so
        # only decode the tail after the pict control word and stop at the
        # closing braces. This is best-effort; linked assets are preferred.
        try:
            data = bytes.fromhex(hex_data)
            img = Image.open(io.BytesIO(data)).convert("RGB")
            images.append(img)
        except Exception:
            continue

    return images


def extract_rtf_images(rtf_path: str):
    """
    Extract source images from RTF in document order.

    Mytonomy's RTF package uses INCLUDEPICTURE fields pointing into Images/.
    The recurring logo.png is ignored. picwgoal/pichgoal are stored in inches,
    matching the HTML style values (e.g. 4320 twips == 3 inches).
    """
    results = []

    raw = Path(rtf_path).read_bytes().decode("latin1", errors="replace")

    # Parse image display dimensions in the same order as INCLUDEPICTURE fields.
    goals = [
        (float(w) / 1440.0, float(h) / 1440.0)
        for w, h in re.findall(
            r"\\pict\b[^{}]*?\\picwgoal(\d+)[^{}]*?\\pichgoal(\d+)",
            raw,
            re.I,
        )
    ]

    refs = _rtf_image_references(rtf_path)

    for idx, ref in enumerate(refs):
        basename = Path(ref).name.lower()

        if basename in IGNORE_IMAGES:
            continue

        asset = _find_rtf_asset(rtf_path, ref)
        if not asset:
            continue

        try:
            img = Image.open(asset).convert("RGB")
        except Exception:
            continue

        # The recurring Mytonomy logo/placeholder is identified by its tiny
        # pixel footprint (same rule used for HTML images), not just by
        # filename — filenames vary, but the placeholder dimensions don't.
        # This was previously only applied to the embedded-\pict fallback
        # path below, so linked INCLUDEPICTURE logo assets (the common case
        # for these RTF packages) were never actually excluded here.
        if _is_logo_placeholder_image(img):
            continue

        if idx < len(goals):
            img._qa_source_display_width = goals[idx][0]
            img._qa_source_display_height = goals[idx][1]

        img._qa_source_name = basename
        results.append(img)

    # If there were no linked assets, try embedded raster pict blocks.
    if not results:
        results = [
            img for img in _extract_embedded_rtf_images(rtf_path)
            if not _is_logo_placeholder_image(img)
        ]

    return results


def extract_rtf_tokens(rtf_path: str):
    """Tokenize RTF text using the same Unicode token model as PDF/HTML."""
    tokens = []
    text = extract_rtf_text(rtf_path)

    for line_no, line in enumerate(text.splitlines(), start=1):
        line = normalize_for_token_scan(line)
        if not line:
            continue

        for m in TOKEN_RE.finditer(line):
            raw = m.group(0)

            if should_ignore_token(raw):
                continue

            tokens.append(
                TextToken(
                    text=raw,
                    norm=norm_token(raw),
                    line=line_no,
                    tag=None,
                    bold=False,
                    italic=False,
                    heading_level=0,
                    in_allowed_block=True,
                )
            )

    return tokens

def is_rtf_source(path: str) -> bool:
    return Path(path).suffix.lower() == ".rtf"


def is_pdf_source(path: str) -> bool:
    return Path(path).suffix.lower() == ".pdf"


def is_docx_target(path: str) -> bool:
    return Path(path).suffix.lower() in (".docx", ".dotx")


# ── DOCX (Word) target support ──────────────────────────────────────────────
#
# PDF ↔ WORD reuses the exact same downstream checks (check_strict_text_mapping,
# check_images, check_garbled, check_placeholders, etc.) that already work for
# HTML targets. To do that cheaply and without duplicating logic, a .docx file
# is converted into a small synthetic HTML document (paragraphs -> <p>, bold
# runs -> <b>, headings -> <h1..h6>, tables -> <table>/<tr>/<td>, inline images
# -> <img> with a data: URI) and parsed with the same BeautifulSoup/html.parser
# pipeline used for real HTML. This keeps qa_engine's core text/image
# comparison logic as the single source of truth for all three modes.

def _docx_heading_tag(paragraph) -> str:
    style_name = (getattr(paragraph.style, "name", "") or "").lower()
    for level in range(1, 7):
        if style_name == f"heading {level}" or style_name == f"heading{level}":
            return f"h{level}"
    if style_name in ("title",):
        return "h1"
    return "p"


def _docx_paragraph_to_html(paragraph) -> str:
    import html as _html_mod

    tag = _docx_heading_tag(paragraph)
    run_html = []
    for run in paragraph.runs:
        text = _html_mod.escape(run.text or "")
        if not text:
            continue
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        if run.underline:
            text = f"<u>{text}</u>"
        run_html.append(text)
    inner = "".join(run_html)
    if not inner.strip():
        return ""
    return f"<{tag}>{inner}</{tag}>"


def _docx_table_to_html(table) -> str:
    import html as _html_mod

    rows_html = []
    for row in table.rows:
        cells_html = []
        for cell in row.cells:
            cell_text = _html_mod.escape(cell.text or "")
            cells_html.append(f"<td>{cell_text}</td>")
        rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
    return f"<table>{''.join(rows_html)}</table>"


def load_docx_as_soup(docx_path: str) -> BeautifulSoup:
    """
    Parse a .docx target into a BeautifulSoup document built from its body
    paragraphs and tables, in document order, so the existing HTML-shaped
    checks (text mapping, garbled text, placeholders, bold-tag checks) work
    unmodified against a Word target.
    """
    try:
        import docx
        from docx.document import Document as _DocxDocument
        from docx.oxml.ns import qn
        from docx.table import Table as _DocxTable
        from docx.text.paragraph import Paragraph as _DocxParagraph
    except Exception:
        return BeautifulSoup("<html><body></body></html>", "html.parser", store_line_numbers=True)

    try:
        document = docx.Document(docx_path)
    except Exception:
        return BeautifulSoup("<html><body></body></html>", "html.parser", store_line_numbers=True)

    def iter_block_items(parent):
        # Walk the document body in true document order (paragraphs + tables interleaved).
        parent_elm = parent.element.body if isinstance(parent, _DocxDocument) else parent._tc
        for child in parent_elm.iterchildren():
            if child.tag == qn("w:p"):
                yield _DocxParagraph(child, parent)
            elif child.tag == qn("w:tbl"):
                yield _DocxTable(child, parent)

    parts = []
    try:
        for block in iter_block_items(document):
            if isinstance(block, _DocxParagraph):
                html_piece = _docx_paragraph_to_html(block)
                if html_piece:
                    parts.append(html_piece)
            elif isinstance(block, _DocxTable):
                parts.append(_docx_table_to_html(block))
    except Exception:
        # Fall back to a flat paragraph walk if the interleaved walk fails
        # for an unusual document structure — better a flattened but non-
        # crashing comparison than none at all.
        for paragraph in document.paragraphs:
            html_piece = _docx_paragraph_to_html(paragraph)
            if html_piece:
                parts.append(html_piece)
        for table in document.tables:
            parts.append(_docx_table_to_html(table))

    # Inline images: represent each as an <img> tag with a data: URI so
    # extract_html_images-style logic (perceptual hashing) can compare them
    # the same way it compares base64 HTML images. The URI MUST include the
    # image subtype (data:image/png;base64,...) -- a bare "data:image;base64,"
    # (missing "/<subtype>") fails the "data:image/" prefix check in the
    # decoder and every docx image was being reported as a broken/undecodable
    # image, verified directly against real source files.
    for img_bytes, img_name in _docx_inline_images(document):
        b64 = base64.b64encode(img_bytes).decode("ascii")
        mime_subtype = "png"
        try:
            with Image.open(io.BytesIO(img_bytes)) as _probe:
                if _probe.format:
                    mime_subtype = _probe.format.lower()
        except Exception:
            ext = Path(img_name).suffix.lstrip(".").lower()
            if ext:
                mime_subtype = "jpeg" if ext in ("jpg", "jpeg") else ext
        parts.append(f'<img src="data:image/{mime_subtype};base64,{b64}" alt="{img_name}">')

    body_html = "\n".join(parts)
    full_html = f"<html><body>{body_html}</body></html>"
    return BeautifulSoup(full_html, "html.parser", store_line_numbers=True)


def _docx_inline_images(document):
    """Yield (image_bytes, part_name) for every image embedded in the docx package."""
    results = []
    try:
        for rel in document.part.rels.values():
            if "image" in rel.reltype:
                try:
                    results.append((rel.target_part.blob, rel.target_part.partname.split("/")[-1]))
                except Exception:
                    continue
    except Exception:
        pass
    return results


def extract_docx_images(docx_path: str, soup: "BeautifulSoup" = None):
    """
    Extract target images from a .docx file, returned in the SAME shape as
    extract_html_images (list of (tag, PIL.Image|None, resolved)) so
    check_images() does not need any Word-specific branching.

    load_docx_as_soup() already embeds every inline docx image as a
    data:image;base64,... <img> tag in document order, so the simplest,
    least-duplicated way to extract them is to hand that soup to the
    existing extract_html_images() data-URI decoder. Each image is opened
    defensively — a single corrupt/unsupported image must not crash the
    whole file's QA (this mirrors the RTF tuple/metadata bug fix below).
    """
    if soup is None:
        soup = load_docx_as_soup(docx_path)
    try:
        return extract_html_images(soup, docx_path)
    except Exception:
        return []


def extract_source_text(source_path: str) -> str:
    if is_rtf_source(source_path):
        return extract_rtf_text(source_path)
    return extract_pdf_text(source_path)


def extract_source_tokens(source_path: str):
    if is_rtf_source(source_path):
        return extract_rtf_tokens(source_path)
    return extract_pdf_tokens(source_path)

def check_strict_text_mapping(pdf_path: str, soup: BeautifulSoup):
    pdf_tokens = merge_spaced_hyphen_tokens(extract_source_tokens(pdf_path))
    html_tokens = merge_spaced_hyphen_tokens(extract_html_tokens(soup))

    pdf_norms = [t.norm for t in pdf_tokens]
    html_norms = [t.norm for t in html_tokens]

    matcher = difflib.SequenceMatcher(
        None,
        pdf_norms,
        html_norms,
        autojunk=False,
    )

    issues = []
    seen = set()
    semantic_seen = set()

    def add_semantic_issue_once(category, severity, html_token, message, expected, actual):
        tag = html_token.tag
        line = html_token.line
        tag_id = id(tag) if tag else None

        key = (
            category,
            line,
            tag_id,
            expected,
            actual,
        )

        if key in semantic_seen:
            return

        semantic_seen.add(key)

        if tag:
            snippet = str(tag)[:300]
            actual_value = actual or tag.name
        else:
            snippet = html_token.text
            actual_value = actual or html_token.text

        _add_issue(
            issues,
            seen,
            Issue(
                category,
                severity,
                line,
                message,
                snippet=snippet,
                expected=expected,
                actual=actual_value,
            ),
        )

    for op, i1, i2, j1, j2 in matcher.get_opcodes():

        if op == "equal":
            for offset in range(i2 - i1):
                p = pdf_tokens[i1 + offset]
                h = html_tokens[j1 + offset]

                # Same normalized token, but different case.
                # Example: Body vs boDy
                if (
                    p.text != h.text
                    and p.text.lower() == h.text.lower()
                    and has_case(p.text)
                    and has_case(h.text)
                ):
                    _add_issue(
                        issues,
                        seen,
                        Issue(
                            "Capitalization Mismatch",
                            "error",
                            h.line,
                            f'Expected capitalization "{p.text}" but HTML shows "{h.text}".',
                            snippet=(
                                f"Source context: {context(pdf_tokens, i1 + offset)}\n"
                                f"HTML context: {context(html_tokens, j1 + offset)}"
                            ),
                            expected=p.text,
                            actual=h.text,
                        ),
                    )

                # Missing bold/strong.
                # Do not require <strong>/<b> for headings because h1-h6 are naturally bold.
                if p.bold and not p.heading_level and not h.bold:
                    add_semantic_issue_once(
                        "Missing Strong/Bold Tag",
                        "warning",
                        h,
                        "Source text is bold, but the matching HTML text is not inside <strong>/<b> or bold CSS.",
                        "strong/b or font-weight:bold",
                        h.tag.name if h.tag else h.text,
                    )

                # Missing heading tag.
                if p.heading_level and not h.heading_level:
                    add_semantic_issue_once(
                        "Missing Heading Tag",
                        "warning",
                        h,
                        f"Source heading text should be inside an h{p.heading_level} heading tag.",
                        f"h{p.heading_level}",
                        h.tag.name if h.tag else h.text,
                    )

                # Heading level mismatch.
                if p.heading_level and h.heading_level and p.heading_level != h.heading_level:
                    add_semantic_issue_once(
                        "Heading Level Mismatch",
                        "warning",
                        h,
                        f"Source heading level is h{p.heading_level}, but HTML uses h{h.heading_level}.",
                        f"h{p.heading_level}",
                        f"h{h.heading_level}",
                    )

                # Text outside expected semantic document blocks.
                # Do not duplicate a missing-heading error with a generic paragraph/list warning.
                if not h.in_allowed_block and not p.heading_level:
                    add_semantic_issue_once(
                        "Missing Paragraph/List Tag",
                        "warning",
                        h,
                        "Visible text is not inside an expected document tag such as p, li, h1-h6, td/th, or figcaption.",
                        "p/li/h1-h6/td/th/figcaption",
                        h.tag.name if h.tag else h.text,
                    )

            continue

        expected = " ".join(t.text for t in pdf_tokens[i1:i2]).strip()
        actual = " ".join(t.text for t in html_tokens[j1:j2]).strip()

        line = html_tokens[j1].line if j1 < len(html_tokens) else None

        if not expected and not actual:
            continue

        # Ignore converter/PDF extraction differences around hyphen spacing.
        # Example:
        #   PDF:  mini - stroke
        #   HTML: mini-stroke
        if _is_hyphen_spacing_noise(expected, actual) or is_hyphen_spacing_only_difference(expected, actual):
            continue

        if op == "replace":
            category = "Changed Text"

            pdf_part = pdf_tokens[i1:i2]
            html_part = html_tokens[j1:j2]

            if pdf_part and html_part:
                if (
                    all(is_punctuation(t.text) for t in pdf_part)
                    or all(is_punctuation(t.text) for t in html_part)
                ):
                    category = "Punctuation Mismatch"

            _add_issue(
                issues,
                seen,
                Issue(
                    category,
                    "error",
                    line,
                    f'Expected "{expected}" but HTML shows "{actual}".',
                    snippet=(
                        f"Source context: {context(pdf_tokens, i1)}\n"
                        f"HTML context: {context(html_tokens, j1)}"
                    ),
                    expected=expected,
                    actual=actual,
                ),
            )

        elif op == "delete":
            category = "Missing Text"

            if all(is_punctuation(t.text) for t in pdf_tokens[i1:i2]):
                category = "Missing Punctuation"

            _add_issue(
                issues,
                seen,
                Issue(
                    category,
                    "error",
                    line,
                    f'Source document text is missing from HTML: "{expected}".',
                    snippet=f"Source context: {context(pdf_tokens, i1)}",
                    expected=expected,
                    actual="",
                ),
            )

        elif op == "insert":
            category = "Extra Text"

            if all(is_punctuation(t.text) for t in html_tokens[j1:j2]):
                category = "Extra Punctuation"

            _add_issue(
                issues,
                seen,
                Issue(
                    category,
                    "error",
                    line,
                    f'HTML contains extra text not found in the source document: "{actual}".',
                    snippet=f"HTML context: {context(html_tokens, j1)}",
                    expected="",
                    actual=actual,
                ),
            )

    return issues
    

def canonical_hyphen_text(text: str) -> str:
    """
    Makes these equivalent:
      long - term
      long- term
      long -term
      long-term
    """
    text = normalize_for_compare(text)
    text = uregex.sub(r"[-‐‒–—−]", "-", text)
    text = uregex.sub(r"(?<=[\p{L}\p{M}\p{N}])\s*-\s*(?=[\p{L}\p{M}\p{N}])", "-", text)
    text = uregex.sub(r"\s+", " ", text)
    return text.strip().lower()


def is_hyphen_spacing_only_difference(expected: str, actual: str) -> bool:
    if not expected or not actual:
        return False

    return canonical_hyphen_text(expected) == canonical_hyphen_text(actual)

def normalize_hyphen_spacing(text):
    text = normalize_for_compare(text)
    text = uregex.sub(r"[-‐‒–—−]", "-", text)
    text = uregex.sub(r"(?<=[\p{L}\p{M}\p{N}])\s*-\s*(?=[\p{L}\p{M}\p{N}])", "-", text)
    text = uregex.sub(r"\s+", " ", text).strip()
    return text.lower()


# ── Other checks ─────────────────────────────────────────────────────────────

def check_filenames(pdf_path, html_path):
    pb = os.path.splitext(os.path.basename(pdf_path))[0]
    hb = os.path.splitext(os.path.basename(html_path))[0]

    if pb != hb:
        return [
            Issue(
                "Filename Mismatch",
                "error",
                None,
                f'Source file is "{os.path.basename(pdf_path)}" but HTML is "{os.path.basename(html_path)}". Basenames should usually match.',
                expected=pb,
                actual=hb,
            )
        ]

    return []


def check_garbled(soup: BeautifulSoup):
    issues = []

    for node in soup.find_all(string=True):
        if not is_visible_text_node(node):
            continue

        text = str(node)

        if MOJIBAKE.search(text):
            line = getattr(node.parent, "sourceline", None)
            issues.append(
                Issue(
                    "Garbled / Encoding Error",
                    "error",
                    line,
                    "Text contains mojibake characters indicating an encoding problem.",
                    snippet=text.strip()[:160],
                    expected="Readable Unicode text",
                    actual=text.strip()[:80],
                )
            )

    return issues


def check_placeholders(soup: BeautifulSoup):
    issues = []

    for node in soup.find_all(string=True):
        if not is_visible_text_node(node):
            continue

        text = str(node)
        m = PLACEHOLDER.search(text)

        if m:
            line = getattr(node.parent, "sourceline", None)
            issues.append(
                Issue(
                    "Leftover Placeholder",
                    "error",
                    line,
                    f'Placeholder text found: "{m.group(0)}".',
                    snippet=text.strip()[:160],
                    expected="Final content",
                    actual=m.group(0),
                )
            )

    return issues


def check_links(soup: BeautifulSoup, html_path: str):
    base = os.path.dirname(os.path.abspath(html_path))
    issues = []

    for a in soup.find_all("a"):
        href = a.get("href", "")

        if not href or href.startswith(("http://", "https://", "mailto:", "#", "javascript:")):
            continue

        resolved = os.path.normpath(os.path.join(base, href))

        if not os.path.exists(resolved):
            line = getattr(a, "sourceline", None)
            issues.append(
                Issue(
                    "Broken Link",
                    "error",
                    line,
                    f'Link target not found: "{href}".',
                    expected="Existing local link target",
                    actual=href,
                )
            )

    return issues

def check_inline_style_rules(soup: BeautifulSoup):
    issues = []
    seen = set()

    def add_style_issue(tag, message):
        line = getattr(tag, "sourceline", None)
        snippet = str(tag)[:250]

        key = (
            "Style Attribute Mismatch",
            line,
            snippet,
            "text-align: left",
            "text-align: right",
        )

        if key in seen:
            return

        seen.add(key)

        issues.append(Issue(
            "Style Attribute Mismatch",
            "error",
            line,
            message,
            snippet=snippet,
            expected="text-align: left",
            actual="text-align: right",
        ))

    # Rule 1: list item text should not be right-aligned
    for li in soup.find_all("li"):
        for tag in li.find_all(["p", "span", "div"], recursive=True):
            text = tag.get_text(" ", strip=True)
            if not text:
                continue

            actual_align = get_css_value(tag, "text-align")

            if actual_align == "right":
                add_style_issue(
                    tag,
                    "List item text is right-aligned, but it should be left-aligned."
                )

    # Rule 2: normal paragraphs should not be right-aligned
    for tag in soup.find_all("p"):
        text = tag.get_text(" ", strip=True)

        if not text:
            continue

        # IMPORTANT:
        # If this <p> is inside <li>, Rule 1 already handled it.
        # Do not report it again as a normal paragraph.
        if tag.find_parent("li"):
            continue

        if tag.find("img") and not text:
            continue

        if is_boilerplate_line(text):
            continue

        actual_align = get_css_value(tag, "text-align")

        if actual_align == "right":
            add_style_issue(
                tag,
                "Paragraph text is right-aligned, but expected document body text alignment is left."
            )

    return issues
   
def check_list_markers(soup: BeautifulSoup):
    issues = []
    seen = set()

    for ul in soup.find_all("ul"):
        ul_style = parse_inline_style(ul.get("style", ""))
        ul_list_style = ul_style.get("list-style-type", "")

        for li in ul.find_all("li", recursive=False):
            line = getattr(li, "sourceline", None)
            text = li.get_text(" ", strip=True)

            if not text:
                continue

            marker = (li.get("data-list-text") or "").strip()
            li_style = parse_inline_style(li.get("style", ""))
            li_list_style = li_style.get("list-style-type", "")

            # Case 1: bullet marker changed
            # Example: data-list-text="°" instead of data-list-text="•"
            if marker and marker != EXPECTED_BULLET_MARKER:
                key = (line, marker, text[:80])

                if key not in seen:
                    seen.add(key)

                    issues.append(Issue(
                        "Bullet Marker Mismatch",
                        "error",
                        line,
                        "List item bullet marker differs from the expected PDF bullet.",
                        snippet=str(li)[:300],
                        expected=EXPECTED_BULLET_MARKER,
                        actual=marker,
                    ))

            # Case 2: bullet marker removed/hidden
            combined_style = (
                (ul.get("style", "") or "")
                + ";"
                + (li.get("style", "") or "")
            ).replace(" ", "").lower()

            bullet_removed = (
                ul_list_style == "none"
                or li_list_style == "none"
                or "list-style:none" in combined_style
                or "list-style-type:none" in combined_style
            )

            if bullet_removed:
                key = (line, "missing-bullet", text[:80])

                if key not in seen:
                    seen.add(key)

                    issues.append(Issue(
                        "Missing Bullet Marker",
                        "error",
                        line,
                        "List item bullet marker appears to be removed or hidden.",
                        snippet=str(li)[:300],
                        expected=EXPECTED_BULLET_MARKER,
                        actual="missing / hidden bullet",
                    ))

    return issues
    

def check_ul_direct_paragraphs(soup: BeautifulSoup):
    """Catch bullet/list text placed directly inside <ul> instead of inside <li>."""
    issues = []
    seen = set()

    for ul in soup.find_all("ul"):
        for child in ul.find_all("p", recursive=False):
            text = child.get_text(" ", strip=True)
            if not text:
                continue
            line = getattr(child, "sourceline", None)
            key = (line, text[:120])
            if key in seen:
                continue
            seen.add(key)
            issues.append(Issue(
                "Missing List Item Tag",
                "error",
                line,
                "List text appears directly inside <ul>; it should be wrapped in <li data-list-text=\"•\">.",
                snippet=str(child)[:300],
                expected='<li data-list-text="•"><p>...</p></li>',
                actual='<ul><p>...</p></ul>',
            ))

    return issues

# ── Orchestrator ─────────────────────────────────────────────────────────────

def run_checks(pdf_path: str, html_path: str, progress_cb=None, validation_mode=None):
    """
    Returns:
      {
        "language": str,
        "issues": list[Issue]
      }

    validation_mode: optional hint from the caller ("pdf", "rtf", or "word").
    This is accepted for interface compatibility with server.py (which always
    passes it) but the engine determines source/target type primarily by
    inspecting the actual files, since that is more reliable than a caller-
    supplied label. validation_mode is currently only used to decide whether
    the *target* should be parsed as HTML or as a DOCX/Word document.
    """
    def prog(pct, label):
        if progress_cb:
            progress_cb(pct, label)

    prog(5, "Extracting source text and parsing target…")

    target_is_word = (
        str(validation_mode or "").lower() in ("word", "docx", "pdf_word", "pdf-word")
        or is_docx_target(html_path)
    )

    if target_is_word:
        soup = load_docx_as_soup(html_path)
        html_imgs = extract_docx_images(html_path, soup=soup)
    else:
        soup = load_html(html_path)
        html_imgs = extract_html_images(soup, html_path)

    source_text = extract_source_text(pdf_path)

    if is_pdf_source(pdf_path):
        try:
            src_imgs = extract_pdf_images(pdf_path)
        except Exception:
            src_imgs = []
    elif is_rtf_source(pdf_path):
        try:
            src_imgs = extract_rtf_images(pdf_path)
        except Exception:
            src_imgs = []
    else:
        src_imgs = []

    language = detect_language(source_text)

    issues = []

    if not target_is_word:
        # Urdu/Arabic metadata checks: lang, dir="rtl", and meaningful title
        # only make sense for an HTML target (docx has no lang/dir markup).
        issues += check_rtl_language_metadata(soup, language)

    prog(35, "Checking title and HTML tag structure…")
    issues += check_filenames(pdf_path, html_path)
    if is_pdf_source(pdf_path) and not target_is_word:
        issues += check_title(pdf_path, soup)

    if not target_is_word:
        issues += check_raw_html_tags(html_path)
        # NOTE: check_tag_attributes_and_styles / check_required_b_tags_by_company_spec /
        # check_missing_b_tags_strict were hardcoded to one fixed "company spec"
        # (padding-left:23pt, text-indent:-10pt, a single bullet glyph, every h1-h3
        # must contain <b>, etc.) applied blindly to every file regardless of what
        # the actual source document contains. That produced hundreds of false
        # "Style Attribute Mismatch"/"Missing B Tag"/"Tag Attribute Mismatch" errors
        # per file that have nothing to do with whether the HTML/RTF/PDF/Word content
        # actually matches its source. Disabled — real content fidelity is enforced
        # by check_strict_text_mapping (word-for-word + semantic) and check_images
        # (visual mapping) below.

    prog(45, "Running strict source-to-HTML text mapping…")
    issues += check_strict_text_mapping(pdf_path, soup)

    prog(72, "Checking images…")
    try:
        issues += check_images(src_imgs, html_imgs)
    except Exception as img_exc:
        # A single bad/unexpected image (e.g. an extractor returning
        # (metadata, image) instead of a plain image) must not crash the
        # whole file's QA — record it as a reviewable issue instead.
        issues.append(Issue(
            "Image Check Error",
            "warning",
            None,
            f"Image comparison could not complete for this file: {img_exc}",
            expected="Images compared successfully",
            actual=str(img_exc),
        ))
    
    IGNORE_ISSUE_CATEGORIES = {
    "Image Display Size Mismatch",
    "Image Aspect Ratio Mismatch",
    "Image File Aspect Ratio Mismatch",
    }

    issues = [
        issue for issue in issues
        if issue.category not in IGNORE_ISSUE_CATEGORIES
    ]

    # Final safety filter: remove false positives that are only hyphen-spacing differences.
    # This catches cases even if they bypassed the central _add_issue helper.
    issues = [
        issue for issue in issues
        if not _should_ignore_hyphen_only_issue(issue)
    ]

    prog(90, "Checking encoding, placeholders, and links…")
    issues += check_garbled(soup)
    issues += check_placeholders(soup)
    if not target_is_word:
        issues += check_links(soup, html_path)

    # Final safety filter after all checks.
    issues = [issue for issue in issues if not _should_ignore_hyphen_only_issue(issue)]

    prog(100, "Done.")

    return {
        "language": language,
        "issues": issues,
    }
   
def parse_inline_style(style: str) -> dict:
    """
    Converts:
      'padding-left: 23pt;text-align: right;'
    into:
      {'padding-left': '23pt', 'text-align': 'right'}
    """
    result = {}

    for part in (style or "").split(";"):
        if ":" not in part:
            continue

        key, value = part.split(":", 1)
        result[key.strip().lower()] = value.strip().lower()

    return result


def get_css_value(tag, prop: str) -> str:
    if not tag:
        return ""

    style = parse_inline_style(tag.get("style", ""))
    return style.get(prop.lower(), "")
    
# ── RTL / Language / Title metadata checks ───────────────────────────────────

RTL_LANG_CODES = {
    "Urdu": "ur",
    "Arabic": "ar",
}

INVALID_TITLE_VALUES = {"", "-", "—", "_", "untitled", "document", "title"}


def _clean_attr(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def _first_heading_text(soup: BeautifulSoup) -> str:
    heading = soup.find(["h1", "h2", "h3"])
    if not heading:
        return ""
    return _clean_attr(heading.get_text(" ", strip=True))


def _html_lang_value(soup: BeautifulSoup) -> str:
    html = soup.find("html")
    body = soup.find("body")

    for tag in (html, body):
        if not tag:
            continue

        lang = (
            tag.get("lang")
            or tag.get("xml:lang")
            or tag.get("data-lang")
            or ""
        )

        lang = _clean_attr(lang).lower()

        if lang:
            return lang

    return ""


def _has_rtl_direction(soup: BeautifulSoup) -> bool:
    html = soup.find("html")
    body = soup.find("body")

    # Check direct dir attributes first.
    for tag in (html, body):
        if tag and _clean_attr(tag.get("dir", "")).lower() == "rtl":
            return True

    # Check any explicit dir="rtl" container.
    for tag in soup.find_all(attrs={"dir": True}):
        if _clean_attr(tag.get("dir", "")).lower() == "rtl":
            return True

    # Check inline style="direction: rtl".
    for tag in soup.find_all(style=True):
        style = tag.get("style", "")
        if re.search(r"direction\s*:\s*rtl", style, flags=re.IGNORECASE):
            return True

    # Check CSS blocks.
    for style_tag in soup.find_all("style"):
        css = style_tag.get_text(" ", strip=True)
        if re.search(r"direction\s*:\s*rtl", css, flags=re.IGNORECASE):
            return True

    return False


def _rtl_actual_direction_summary(soup: BeautifulSoup) -> str:
    html = soup.find("html")
    body = soup.find("body")

    html_dir = _clean_attr(html.get("dir", "")) if html else ""
    body_dir = _clean_attr(body.get("dir", "")) if body else ""

    if html_dir or body_dir:
        return f'html dir="{html_dir}", body dir="{body_dir}"'

    return "No dir=\"rtl\" found on html/body or text containers"


def check_rtl_language_metadata(soup: BeautifulSoup, detected_language: str):
    """
    Adds missing metadata checks for RTL languages.

    For Urdu/Arabic HTML, the output should usually include:
      <html lang="ur" dir="rtl">
    or:
      <html lang="ar" dir="rtl">

    Also catches placeholder titles like:
      <title>-</title>
    """
    issues = []

    detected_language = detected_language or ""
    expected_lang = RTL_LANG_CODES.get(detected_language)

    if not expected_lang:
        return issues

    html = soup.find("html")
    title_tag = soup.find("title")
    first_heading = _first_heading_text(soup)

    html_snippet = str(html)[:350] if html else "<html> tag not found"

    # 1. Missing / wrong lang attribute
    actual_lang = _html_lang_value(soup)
    actual_lang_base = actual_lang.split("-")[0] if actual_lang else ""

    if actual_lang_base != expected_lang:
        issues.append(Issue(
            category="Missing Language Attribute",
            severity="warning",
            line=getattr(html, "sourceline", None) if html else None,
            message=(
                f'HTML language attribute should be set for {detected_language}. '
                f'Expected lang="{expected_lang}".'
            ),
            snippet=html_snippet,
            expected=f'<html lang="{expected_lang}">',
            actual=f'lang="{actual_lang}"' if actual_lang else "No lang attribute found",
        ))

    # 2. Missing RTL direction
    if not _has_rtl_direction(soup):
        issues.append(Issue(
            category="Missing RTL Direction",
            severity="warning",
            line=getattr(html, "sourceline", None) if html else None,
            message=(
                f'{detected_language} is a right-to-left language, but the HTML '
                f'does not define RTL direction.'
            ),
            snippet=html_snippet,
            expected='<html dir="rtl">',
            actual=_rtl_actual_direction_summary(soup),
        ))

    # 3. Missing / placeholder title
    title_text = _clean_attr(title_tag.get_text(" ", strip=True)) if title_tag else ""
    normalized_title = title_text.lower()

    if normalized_title in INVALID_TITLE_VALUES:
        expected_title = first_heading or f"Meaningful {detected_language} document title"

        issues.append(Issue(
            category="Missing Title",
            severity="error",
            line=getattr(title_tag, "sourceline", None) if title_tag else None,
            message=(
                "HTML title is missing, empty, or only a placeholder. "
                "The title should describe the document content."
            ),
            snippet=str(title_tag)[:350] if title_tag else "<title> tag not found",
            expected=f"<title>{expected_title}</title>",
            actual=f"<title>{title_text}</title>" if title_tag else "No <title> tag found",
        ))

    return issues